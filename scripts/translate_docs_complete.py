#!/usr/bin/env python3
"""
完整文档翻译脚本 - 超级优化版
✅ 保持Word格式不变
✅ 批量翻译（10个段落/次）
✅ 移除页眉页脚
✅ 翻译文件夹名和文件名
✅ 增量处理（跳过已完成）
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import List, Dict
import requests

# 强制刷新输出
sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# ============ 配置 ============
SOURCE_DIR = "docs/specs/Full Length"
OUTPUT_DIR = "docs/specs_zh/Full Length"
TRANSLATION_MAP = "docs/specs/文件名翻译对照表_完整版.json"
VLLM_URL = os.getenv("VLLM_URL", "http://10.10.18.3:8000")
VLLM_MODEL = os.getenv("VLLM_MODEL", "/mnt/data/models/Qwen3-32B")
BATCH_SIZE = 10  # 每批翻译10个段落

# ============ 加载翻译对照表 ============
def load_translation_map() -> Dict:
    """加载文件名翻译对照表"""
    try:
        with open(TRANSLATION_MAP, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {"folders": {}, "files": {}}

TRANS_MAP = load_translation_map()

# ============ 文件名翻译 ============
def translate_folder_name(folder_name: str) -> str:
    """翻译文件夹名（使用预设对照表）"""
    return TRANS_MAP.get("folders", {}).get(folder_name, folder_name)

def translate_file_name(file_name: str) -> str:
    """翻译文件名（保留扩展名）"""
    name_without_ext = file_name.replace('.DOC', '').replace('.docx', '')
    translated = TRANS_MAP.get("files", {}).get(file_name, name_without_ext)
    return f"{translated}.docx"

# ============ 内容过滤 ============
def should_skip_paragraph(text: str) -> bool:
    """判断段落是否应该被跳过（删除）"""
    skip_keywords = [
        "Copyright",
        "The American Institute of Architects",
        "AIA",
        "Exclusively published and distributed by Deltek",
        "TIPS:",
        "To view non-printing Editor's Notes",
        "MasterWorks/Single-File Formatting",
        "MasterWorks/Supporting Information",
        "Content Requests:",
        "<Double click here to submit",
        "Double click here",
    ]

    # 检查是否包含任何跳过关键词
    for keyword in skip_keywords:
        if keyword in text:
            return True

    return False

# ============ 翻译API ============
def translate_single(text: str) -> str:
    """翻译单个文本"""
    if not text.strip():
        return text

    prompt = f"""翻译成中文（保留专业术语和格式）：

{text}

中文："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 1000
            },
            timeout=60
        )

        if response.status_code == 200:
            result = response.json()['choices'][0]['message']['content'].strip()

            # 移除<think>标签内容
            if '<think>' in result:
                if '</think>' in result:
                    result = result.split('</think>')[-1].strip()
                else:
                    # 如果没有闭合标签，返回原文
                    return text

            # 清理前缀
            if result.startswith("中文：") or result.startswith("翻译："):
                result = result.split("：", 1)[1].strip()

            # 只取第一行（避免多余解释）
            result = result.split('\n')[0].strip()

            return result if result else text
        return text
    except:
        return text

def translate_batch(texts: List[str]) -> List[str]:
    """批量翻译多个段落（10个/次）"""
    if not texts:
        return []

    # 过滤空段落
    non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
    non_empty_texts = [texts[i] for i in non_empty_indices]

    if not non_empty_texts:
        return texts

    # 合并文本
    combined = "\n---SPLIT---\n".join(non_empty_texts)

    prompt = f"""你是建筑工程规范翻译专家。翻译以下{len(non_empty_texts)}个段落（用---SPLIT---分隔）：

{combined}

要求：
1. 保留所有专业术语、编号、格式标记
2. 保持原有的换行和缩进
3. 用---SPLIT---分隔翻译结果
4. 直接输出翻译，不要加"翻译："等前缀

中文翻译："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": len(combined) * 3
            },
            timeout=180
        )

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            # 移除<think>标签内容
            if '<think>' in translated:
                if '</think>' in translated:
                    translated = translated.split('</think>')[-1].strip()
                else:
                    # 如果没有闭合，回退到逐个翻译
                    return [translate_single(t) for t in texts]

            # 清理可能的前缀
            if translated.startswith("翻译：") or translated.startswith("中文翻译："):
                translated = translated.split("：", 1)[1].strip()

            # 分割结果
            parts = translated.split("---SPLIT---")

            if len(parts) == len(non_empty_texts):
                # 恢复到原始位置
                results = texts.copy()
                for i, idx in enumerate(non_empty_indices):
                    results[idx] = parts[i].strip()
                return results
            else:
                print(f"  ⚠️  批量分割失败({len(parts)}!={len(non_empty_texts)})，逐个翻译...")
                # 改为逐个翻译
                return [translate_single(t) for t in texts]
        else:
            print(f"  ⚠️  API错误: {response.status_code}")
            return texts
    except Exception as e:
        print(f"  ⚠️  翻译失败: {e}")
        return texts

# ============ Word文档处理 ============
def process_word_document(input_file: Path, output_file: Path):
    """处理Word文档：翻译内容，移除页眉页脚，保持格式"""
    print(f"\n{'='*70}")
    print(f"📄 处理: {input_file.name}")
    print(f"{'='*70}")

    # 1. DOC转DOCX
    if input_file.suffix.upper() == '.DOC':
        temp_docx = Path(f"/tmp/{input_file.stem}_temp.docx")
        if sys.platform == 'darwin':
            print("🔄 DOC转DOCX...")
            result = subprocess.run([
                'textutil', '-convert', 'docx',
                str(input_file), '-output', str(temp_docx)
            ], capture_output=True, text=True)

            if result.returncode != 0 or not temp_docx.exists():
                print(f"❌ 转换失败: {result.stderr}")
                return False
            working_file = temp_docx
        else:
            print("❌ 非macOS系统，跳过")
            return False
    else:
        working_file = input_file

    # 2. 创建输出目录
    output_file.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document

        doc = Document(str(working_file))

        # 3. 移除页眉页脚
        print("🗑️  移除页眉页脚...")
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for p in section.header.paragraphs:
                p.clear()
            for p in section.footer.paragraphs:
                p.clear()

        # 4. 移除文本框和形状
        print("🗑️  移除文本框和形状...")
        try:
            # 移除所有shape（包括文本框）
            for shape in doc.inline_shapes:
                shape._element.getparent().remove(shape._element)
        except:
            pass  # 如果没有shape，忽略错误

        # 5. 过滤并收集需要翻译的段落
        paragraphs_data = []
        skipped_count = 0
        for p in doc.paragraphs:
            if p.text.strip():
                # 检查是否应该跳过（删除）
                if should_skip_paragraph(p.text):
                    p.clear()  # 直接删除该段落
                    skipped_count += 1
                    continue

                # 保存完整格式（包括段落级别和run级别）
                formats = []
                for run in p.runs:
                    formats.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })

                # 保存段落级别格式
                para_format = {
                    'style': p.style,  # 段落样式（PR1, PR2等）
                    'alignment': p.alignment,  # 对齐方式
                    'left_indent': p.paragraph_format.left_indent,  # 左缩进
                    'first_line_indent': p.paragraph_format.first_line_indent,  # 首行缩进
                    'space_before': p.paragraph_format.space_before,  # 段前间距
                    'space_after': p.paragraph_format.space_after,  # 段后间距
                }

                paragraphs_data.append({
                    'paragraph': p,
                    'text': p.text,
                    'formats': formats,
                    'para_format': para_format
                })

        if skipped_count > 0:
            print(f"🗑️  已删除 {skipped_count} 个版权/提示段落")

        total = len(paragraphs_data)
        print(f"📝 批量翻译 {total} 个段落（{BATCH_SIZE}个/批）...")

        # 6. 批量翻译段落
        success_count = 0
        for i in range(0, total, BATCH_SIZE):
            batch = paragraphs_data[i:i+BATCH_SIZE]
            batch_texts = [item['text'] for item in batch]

            progress = f"[{i+1}-{min(i+BATCH_SIZE, total)}/{total}]"
            print(f"{progress} 翻译中...", end='', flush=True)

            translated_texts = translate_batch(batch_texts)

            # 应用翻译（保持格式）
            for j, (item, trans_text) in enumerate(zip(batch, translated_texts)):
                para = item['paragraph']
                para_fmt = item['para_format']

                # 最安全的方法：只修改文本，不动段落结构
                # 这样可以保留段落样式（PR1, ART等）和自动编号（A. B. C.）

                if para.runs:
                    # 如果有多个runs，合并到第一个run
                    if len(para.runs) > 1:
                        # 保存第一个run的格式
                        first_run = para.runs[0]

                        # 删除其他runs（从后往前删）
                        for i in range(len(para.runs) - 1, 0, -1):
                            p = para._element
                            p.remove(para.runs[i]._element)

                        # 修改第一个run的文本
                        first_run.text = trans_text
                    else:
                        # 只有一个run，直接修改
                        para.runs[0].text = trans_text

                    # 应用格式到第一个run
                    if item['formats']:
                        run = para.runs[0]
                        fmt = item['formats'][0]
                        if fmt['bold'] is not None:
                            run.bold = fmt['bold']
                        if fmt['italic'] is not None:
                            run.italic = fmt['italic']
                        if fmt['underline'] is not None:
                            run.underline = fmt['underline']
                        if fmt['font_name']:
                            run.font.name = fmt['font_name']
                        if fmt['font_size']:
                            run.font.size = fmt['font_size']
                        if fmt['font_color']:
                            run.font.color.rgb = fmt['font_color']
                else:
                    # 如果没有run，创建一个
                    run = para.add_run(trans_text)
                    if item['formats']:
                        fmt = item['formats'][0]
                        if fmt['bold'] is not None:
                            run.bold = fmt['bold']
                        if fmt['italic'] is not None:
                            run.italic = fmt['italic']
                        if fmt['underline'] is not None:
                            run.underline = fmt['underline']
                        if fmt['font_name']:
                            run.font.name = fmt['font_name']
                        if fmt['font_size']:
                            run.font.size = fmt['font_size']
                        if fmt['font_color']:
                            run.font.color.rgb = fmt['font_color']

            success_count += len(batch)
            print(f" ✅")

        # 7. 翻译表格
        if doc.tables:
            print(f"\n📊 翻译 {len(doc.tables)} 个表格...")
            for table_idx, table in enumerate(doc.tables, 1):
                cells_data = []
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                cells_data.append({
                                    'paragraph': para,
                                    'text': para.text,
                                    'bold': para.runs[0].bold if para.runs else None,
                                    'italic': para.runs[0].italic if para.runs else None
                                })

                if cells_data:
                    cell_texts = [c['text'] for c in cells_data]
                    print(f"  表格{table_idx}: {len(cell_texts)}个单元格...", end='', flush=True)
                    translated = translate_batch(cell_texts)

                    for cell_data, trans in zip(cells_data, translated):
                        para = cell_data['paragraph']

                        # 安全修改：只修改文本，不清空段落（保留样式和编号）
                        if para.runs:
                            # 删除多余的runs（从后往前）
                            if len(para.runs) > 1:
                                for i in range(len(para.runs) - 1, 0, -1):
                                    p = para._element
                                    p.remove(para.runs[i]._element)
                            # 修改第一个run的文本
                            para.runs[0].text = trans
                            para.runs[0].bold = cell_data['bold']
                            para.runs[0].italic = cell_data['italic']
                        else:
                            run = para.add_run(trans)
                            run.bold = cell_data['bold']
                            run.italic = cell_data['italic']

                    print(" ✅")

        # 8. 保存
        doc.save(str(output_file))
        try:
            rel_path = output_file.relative_to(Path.cwd())
            print(f"\n✅ 完成！保存到: {rel_path}")
        except:
            print(f"\n✅ 完成！保存到: {output_file}")

        # 9. 清理临时文件
        if input_file.suffix.upper() == '.DOC' and temp_docx.exists():
            temp_docx.unlink()

        return True

    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False

# ============ 批量翻译 ============
def batch_translate_all(max_files=None):
    """批量翻译所有文档（带文件名翻译）"""
    source_path = Path(SOURCE_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    # 获取所有子文件夹
    folders = sorted([d for d in source_path.iterdir() if d.is_dir() and not d.name.startswith('.')])

    print(f"\n{'='*70}")
    print(f"  🚀 批量文档翻译（完整版）")
    print(f"{'='*70}")
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"VLLM: {VLLM_URL}")
    print(f"模型: {VLLM_MODEL}")
    print(f"批量: {BATCH_SIZE} 段/次")
    print(f"文件夹: {len(folders)} 个")
    print(f"翻译表: {len(TRANS_MAP.get('folders', {}))} 个文件夹名")
    print(f"{'='*70}\n")

    total_success = 0
    total_fail = 0
    total_skip = 0

    # 遍历所有文件夹
    for folder_idx, folder in enumerate(folders, 1):
        folder_name = folder.name
        translated_folder_name = translate_folder_name(folder_name)

        print(f"\n{'#'*70}")
        print(f"# 文件夹 [{folder_idx}/{len(folders)}]: {folder_name}")
        print(f"# 翻译为: {translated_folder_name}")
        print(f"{'#'*70}")

        # 获取该文件夹下的所有文件
        files = sorted(list(folder.glob("*.DOC")) + list(folder.glob("*.docx")))
        # 过滤临时文件：~$ 开头的和 .~ 开头的
        files = [f for f in files if not f.name.startswith('~$') and not f.name.startswith('.~')]

        if max_files:
            files = files[:max_files]

        # 处理每个文件
        for file_idx, input_file in enumerate(files, 1):
            file_name = input_file.name
            translated_file_name = translate_file_name(file_name)

            # 输出路径
            output_folder = Path(OUTPUT_DIR).parent / translated_folder_name
            output_file = output_folder / translated_file_name

            print(f"\n  [{file_idx}/{len(files)}] {file_name}")
            print(f"       → {translated_file_name}")

            # 检查是否已存在
            if output_file.exists():
                print(f"  ⏭️  已存在，跳过")
                total_skip += 1
                continue

            # 翻译文档
            try:
                success = process_word_document(input_file, output_file)
                if success:
                    total_success += 1
                else:
                    total_fail += 1
            except KeyboardInterrupt:
                print("\n\n⚠️  用户中断")
                break
            except Exception as e:
                print(f"\n  ❌ 失败: {e}")
                total_fail += 1

        if max_files:
            break  # 测试模式只处理一个文件夹

    # 统计
    print(f"\n\n{'='*70}")
    print(f"  📊 翻译完成")
    print(f"{'='*70}")
    print(f"✅ 成功: {total_success}")
    print(f"⏭️  跳过: {total_skip}")
    print(f"❌ 失败: {total_fail}")
    print(f"输出: {OUTPUT_DIR}")
    print(f"{'='*70}\n")

# ============ 主程序 ============
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="完整文档翻译（优化版）")
    parser.add_argument("--test", action="store_true", help="测试模式（1个文件）")
    parser.add_argument("--small", action="store_true", help="小批量（10个文件）")

    args = parser.parse_args()

    # 测试VLLM连接
    print("🔌 测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM服务连接成功: {VLLM_URL}\n")
        else:
            print(f"⚠️  VLLM响应异常: {response.status_code}")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        print("请检查VLLM服务是否运行")
        sys.exit(1)

    max_files = None
    if args.test:
        max_files = 1
    elif args.small:
        max_files = 10

    batch_translate_all(max_files)
