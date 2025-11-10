#!/usr/bin/env python3
"""
测试07目录的第一个文件
"""

import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from batch_convert_and_translate import (
    convert_doc_to_docx,
    translate_docx,
    translate_file_name,
    translate_folder_name,
    VLLM_URL
)
import requests

# 测试文件
TEST_FILE = "docs/specs/Full Length/07 - THERMAL AND MOISTURE PROTECTION/070150.19 FL - Preparation for Reroofing.DOC"
FOLDER_NAME = "07 - THERMAL AND MOISTURE PROTECTION"

def test_07_first():
    print("\n" + "="*70)
    print("  🧪 测试07目录第一个文件")
    print("="*70)

    input_file = Path(TEST_FILE)
    if not input_file.exists():
        print(f"❌ 文件不存在: {TEST_FILE}")
        return

    # 翻译文件夹名
    translated_folder = translate_folder_name(FOLDER_NAME)
    print(f"\n📁 文件夹: {FOLDER_NAME}")
    print(f"   翻译为: {translated_folder}")

    # 翻译文件名
    print(f"\n📝 原文件名: {input_file.name}")
    translated_name = translate_file_name(input_file.name)
    print(f"   翻译后: {translated_name}")

    # 输出文件
    output_folder = Path(f"docs/specs_zh/{translated_folder}")
    output_file = output_folder / translated_name

    if output_file.exists():
        print(f"\n🗑️  删除旧文件: {output_file}")
        output_file.unlink()

    # 步骤1：转换
    print(f"\n🔄 步骤1: DOC → DOCX 转换...")
    if convert_doc_to_docx(input_file, output_file):
        print(f"   ✅ 转换成功")
    else:
        print(f"   ❌ 转换失败")
        return

    # 步骤2：翻译（并设置宋体）
    print(f"\n📝 步骤2: 翻译内容并设置宋体...")
    if translate_docx(output_file):
        print(f"   ✅ 翻译成功")
    else:
        print(f"   ❌ 翻译失败")
        return

    # 步骤3：验证
    print(f"\n🔍 步骤3: 验证结果...")
    verify_result(output_file)

    print(f"\n✅ 测试完成！")
    print(f"   输出文件: {output_file}")

def verify_result(file_path):
    """验证翻译结果和字体"""
    from docx import Document

    doc = Document(str(file_path))

    total_paras = len(doc.paragraphs)
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]

    # 统计样式
    styles = {}
    for para in doc.paragraphs:
        style = para.style.name
        if style not in styles:
            styles[style] = 0
        styles[style] += 1

    # 检查字体
    fonts = {}
    for para in non_empty_paras[:20]:  # 检查前20个段落
        for run in para.runs:
            if run.text.strip():
                font_name = run.font.name
                if font_name:
                    fonts[font_name] = fonts.get(font_name, 0) + 1

    # 检查英文残留
    english_keywords = ["copyright", "double click", "tips:", "aia", "<insert"]
    english_paras = []
    for para in non_empty_paras:
        text = para.text.lower()
        if any(keyword in text for keyword in english_keywords):
            english_paras.append(para.text[:80])

    print(f"   总段落数: {total_paras}")
    print(f"   非空段落: {len(non_empty_paras)}")
    print(f"   样式分布: {styles}")
    print(f"   字体分布: {fonts}")

    if "宋体" in fonts or "SimSun" in fonts:
        print(f"   ✅ 宋体设置成功")
    else:
        print(f"   ⚠️  未检测到宋体")

    if english_paras:
        print(f"   ⚠️  英文残留:")
        for text in english_paras[:3]:
            print(f"      - {text}")
    else:
        print(f"   ✅ 无英文残留")

if __name__ == "__main__":
    # 测试VLLM连接
    print("🔌 测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM连接成功\n")
        else:
            print(f"⚠️  VLLM响应异常")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        sys.exit(1)

    test_07_first()
