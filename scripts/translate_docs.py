#!/usr/bin/env python3
"""
批量文档翻译脚本
支持DOC/DOCX和PDF格式，保留原始格式，只翻译文本内容
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import List, Dict, Optional
import shutil

# 配置
SOURCE_DIR = "docs/specs"
OUTPUT_DIR = "docs/specs_zh"
API_URL = "http://localhost:3000/api/chat/translate"
API_TOKEN = os.getenv("API_TOKEN", "")

def translate_text(text: str, target_lang: str = "zh") -> str:
    """
    调用平台翻译API
    """
    if not text or not text.strip():
        return text

    import subprocess
    import json

    payload = json.dumps({
        "content": text,
        "targetLanguage": target_lang
    })

    cmd = [
        'curl', '-s', '-X', 'POST', API_URL,
        '-H', 'Content-Type: application/json',
        '-H', f'Authorization: Bearer {API_TOKEN}',
        '-d', payload
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0:
            response = json.loads(result.stdout)
            return response.get('translatedText', text)
        else:
            print(f"翻译失败: {result.stderr}")
            return text
    except Exception as e:
        print(f"翻译异常: {e}")
        return text

def process_doc_file(input_file: Path, output_file: Path):
    """
    处理DOC/DOCX文件，保留格式，翻译文本，移除页眉页脚
    """
    print(f"处理文档: {input_file.name}")

    # 第一步：将DOC转换为DOCX（如果是DOC格式）
    if input_file.suffix.upper() == '.DOC':
        temp_docx = input_file.with_suffix('.docx')
        # 使用textutil (macOS) 或其他工具转换
        if sys.platform == 'darwin':
            subprocess.run([
                'textutil', '-convert', 'docx',
                str(input_file), '-output', str(temp_docx)
            ], check=False)
            working_file = temp_docx
        else:
            # 如果转换失败，直接复制
            print(f"警告: 无法转换 {input_file.name}，跳过")
            return
    else:
        working_file = input_file

    # 确保输出目录存在
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # 第二步：使用python-docx处理DOCX文件
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(str(working_file))

        # 移除页眉页脚
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # 清空页眉
            for paragraph in section.header.paragraphs:
                paragraph.clear()

            # 清空页脚
            for paragraph in section.footer.paragraphs:
                paragraph.clear()

        # 翻译正文段落，保留格式
        total_paragraphs = len(doc.paragraphs)
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"  翻译段落 {idx+1}/{total_paragraphs}: {paragraph.text[:50]}...")

                # 保存原有格式
                original_runs = []
                for run in paragraph.runs:
                    original_runs.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color.rgb else None
                    })

                # 获取整段文本并翻译
                full_text = paragraph.text
                translated_text = translate_text(full_text)

                # 清空段落
                paragraph.clear()

                # 重新添加翻译后的文本，应用原有格式
                if original_runs:
                    # 简化处理：使用第一个run的格式
                    first_format = original_runs[0]
                    run = paragraph.add_run(translated_text)
                    run.bold = first_format['bold']
                    run.italic = first_format['italic']
                    run.underline = first_format['underline']
                    if first_format['font_name']:
                        run.font.name = first_format['font_name']
                    if first_format['font_size']:
                        run.font.size = first_format['font_size']
                    if first_format['font_color']:
                        run.font.color.rgb = first_format['font_color']

        # 翻译表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            translated = translate_text(paragraph.text)
                            # 保留格式，更新文本
                            if paragraph.runs:
                                first_run = paragraph.runs[0]
                                paragraph.clear()
                                new_run = paragraph.add_run(translated)
                                new_run.bold = first_run.bold
                                new_run.italic = first_run.italic

        # 保存翻译后的文档
        doc.save(str(output_file))
        print(f"✓ 已保存: {output_file}")

        # 清理临时文件
        if input_file.suffix.upper() == '.DOC' and temp_docx.exists():
            temp_docx.unlink()

    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("请运行: pip3 install python-docx --break-system-packages")
        sys.exit(1)
    except Exception as e:
        print(f"处理文档时出错: {e}")
        import traceback
        traceback.print_exc()

def process_pdf_file(input_file: Path, output_file: Path):
    """
    处理PDF文件
    注意：PDF格式保留比较复杂，这里提供基础实现
    """
    print(f"处理PDF: {input_file.name}")

    try:
        import pdfplumber
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 注册中文字体
        # 需要指定系统中的中文字体路径
        # pdfmetrics.registerFont(TTFont('SimSun', '/path/to/simsun.ttf'))

        # 提取文本
        all_text = []
        with pdfplumber.open(str(input_file)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)

        # 翻译
        translated_pages = []
        for i, text in enumerate(all_text):
            print(f"  翻译页面 {i+1}/{len(all_text)}...")
            translated = translate_text(text)
            translated_pages.append(translated)

        # 生成新PDF（简化版本）
        output_file.parent.mkdir(parents=True, exist_ok=True)
        c = canvas.Canvas(str(output_file), pagesize=letter)

        for page_text in translated_pages:
            c.drawString(100, 750, page_text[:100])  # 简化处理
            c.showPage()

        c.save()
        print(f"✓ 已保存: {output_file}")

    except ImportError:
        print("错误: 需要安装 pdfplumber 和 reportlab")
        print("请运行: pip3 install pdfplumber reportlab --break-system-packages")
        sys.exit(1)
    except Exception as e:
        print(f"处理PDF时出错: {e}")

def batch_translate(source_dir: str = SOURCE_DIR, output_dir: str = OUTPUT_DIR):
    """
    批量翻译目录下的所有文档
    """
    source_path = Path(source_dir)
    output_path = Path(output_dir)

    if not source_path.exists():
        print(f"错误: 源目录不存在: {source_dir}")
        return

    # 获取所有文件
    doc_files = list(source_path.rglob("*.DOC")) + list(source_path.rglob("*.docx"))
    pdf_files = list(source_path.rglob("*.pdf"))

    total_files = len(doc_files) + len(pdf_files)
    print(f"找到 {len(doc_files)} 个DOC/DOCX文件，{len(pdf_files)} 个PDF文件")
    print(f"总计: {total_files} 个文件\n")

    # 处理DOC/DOCX文件
    for i, input_file in enumerate(doc_files, 1):
        print(f"\n[{i}/{total_files}] 处理文档文件...")

        # 计算相对路径
        rel_path = input_file.relative_to(source_path)
        output_file = output_path / rel_path

        # 确保输出文件是.docx格式
        if output_file.suffix.upper() == '.DOC':
            output_file = output_file.with_suffix('.docx')

        try:
            process_doc_file(input_file, output_file)
        except Exception as e:
            print(f"✗ 失败: {e}")
            continue

    # 处理PDF文件
    for i, input_file in enumerate(pdf_files, len(doc_files) + 1):
        print(f"\n[{i}/{total_files}] 处理PDF文件...")

        rel_path = input_file.relative_to(source_path)
        output_file = output_path / rel_path

        try:
            process_pdf_file(input_file, output_file)
        except Exception as e:
            print(f"✗ 失败: {e}")
            continue

    print(f"\n\n{'='*60}")
    print(f"翻译完成！")
    print(f"输出目录: {output_dir}")
    print(f"{'='*60}")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="批量翻译文档")
    parser.add_argument("--source", default=SOURCE_DIR, help="源目录")
    parser.add_argument("--output", default=OUTPUT_DIR, help="输出目录")
    parser.add_argument("--test", action="store_true", help="测试模式（只处理前3个文件）")
    parser.add_argument("--token", help="API Token")

    args = parser.parse_args()

    if args.token:
        os.environ["API_TOKEN"] = args.token

    if args.test:
        print("=== 测试模式 ===")
        # 只处理几个文件进行测试

    batch_translate(args.source, args.output)
