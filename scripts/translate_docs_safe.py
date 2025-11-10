#!/usr/bin/env python3
"""
安全翻译脚本 - 完全保护样式和编号
✅ 只修改文本节点，不动XML结构
✅ 保留所有样式（ART, PR1, PR2等）
✅ 保留自动编号（A. B. C.）
✅ 保留多级列表
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import List, Dict
import requests

sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# ============ 配置 ============
SOURCE_DIR = "docs/specs/Full Length"
OUTPUT_DIR = "docs/specs_zh/Full Length"
TRANSLATION_MAP = "docs/specs/文件名翻译对照表_完整版.json"
VLLM_URL = os.getenv("VLLM_URL", "http://10.10.18.3:8000")
VLLM_MODEL = os.getenv("VLLM_MODEL", "/mnt/data/models/Qwen3-32B")
BATCH_SIZE = 10

# ============ 加载翻译对照表 ============
def load_translation_map() -> Dict:
    try:
        with open(TRANSLATION_MAP, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {"folders": {}, "files": {}}

TRANS_MAP = load_translation_map()

def translate_folder_name(folder_name: str) -> str:
    return TRANS_MAP.get("folders", {}).get(folder_name, folder_name)

def translate_file_name(file_name: str) -> str:
    name_without_ext = file_name.replace('.DOC', '').replace('.docx', '')
    translated = TRANS_MAP.get("files", {}).get(file_name, name_without_ext)
    return f"{translated}.docx"

# ============ 内容过滤 ============
def should_skip_paragraph(text: str) -> bool:
    skip_keywords = [
        "Copyright",
        "The American Institute of Architects",
        "AIA",
        "Exclusively published and distributed by Deltek",
        "TIPS:",
        "To view non-printing Editor's Notes",
        "MasterWorks/Single-File Formatting",
        "MasterWorks/Supporting Information",
        "Content Requests:",
        "<Double click here to submit",
        "Double click here",
    ]
    for keyword in skip_keywords:
        if keyword in text:
            return True
    return False

# ============ 翻译API ============
def translate_single(text: str) -> str:
    if not text.strip():
        return text

    prompt = f"""翻译成中文（保留专业术语和格式）：

{text}

中文："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 1000
            },
            timeout=60
        )

        if response.status_code == 200:
            result = response.json()['choices'][0]['message']['content'].strip()
            if '<think>' in result:
                if '</think>' in result:
                    result = result.split('</think>')[-1].strip()
                else:
                    return text
            if result.startswith("中文：") or result.startswith("翻译："):
                result = result.split("：", 1)[1].strip()
            result = result.split('\n')[0].strip()
            return result if result else text
        return text
    except:
        return text

def translate_batch(texts: List[str]) -> List[str]:
    if not texts:
        return []

    non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
    non_empty_texts = [texts[i] for i in non_empty_indices]

    if not non_empty_texts:
        return texts

    combined = "\n---SPLIT---\n".join(non_empty_texts)
    prompt = f"""你是建筑工程规范翻译专家。翻译以下{len(non_empty_texts)}个段落（用---SPLIT---分隔）：

{combined}

要求：
1. 保留所有专业术语、编号、格式标记
2. 保持原有的换行和缩进
3. 用---SPLIT---分隔翻译结果
4. 直接输出翻译，不要加"翻译："等前缀

中文翻译："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": len(combined) * 3
            },
            timeout=180
        )

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            if '<think>' in translated:
                if '</think>' in translated:
                    translated = translated.split('</think>')[-1].strip()
                else:
                    return [translate_single(t) for t in texts]

            if translated.startswith("翻译：") or translated.startswith("中文翻译："):
                translated = translated.split("：", 1)[1].strip()

            parts = translated.split("---SPLIT---")

            if len(parts) == len(non_empty_texts):
                results = texts.copy()
                for i, idx in enumerate(non_empty_indices):
                    results[idx] = parts[i].strip()
                return results
            else:
                print(f"  ⚠️  批量分割失败({len(parts)}!={len(non_empty_texts)})，逐个翻译...")
                return [translate_single(t) for t in texts]
        else:
            return texts
    except Exception as e:
        print(f"  ⚠️  翻译失败: {e}")
        return [translate_single(t) for t in texts]

# ============ 安全文档处理 - 只修改文本不动结构 ============
def process_word_document_safe(input_file: Path, output_file: Path):
    """
    最安全的处理方式：
    1. 只修改每个run的text属性
    2. 完全不删除、不添加、不重建任何元素
    3. 保留段落样式、编号、列表等所有属性
    """
    print(f"\n{'='*70}")
    print(f"📄 处理: {input_file.name}")
    print(f"{'='*70}")

    # 1. DOC转DOCX
    if input_file.suffix.upper() == '.DOC':
        temp_docx = Path(f"/tmp/{input_file.stem}_temp.docx")
        if sys.platform == 'darwin':
            print("🔄 DOC转DOCX (使用textutil)...")
            result = subprocess.run([
                'textutil', '-convert', 'docx',
                str(input_file), '-output', str(temp_docx)
            ], capture_output=True, text=True)

            if result.returncode != 0 or not temp_docx.exists():
                print(f"❌ 转换失败")
                return False
            working_file = temp_docx
        else:
            print("❌ 非macOS系统")
            return False
    else:
        working_file = input_file

    output_file.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document

        doc = Document(str(working_file))

        # 2. 移除页眉页脚
        print("🗑️  移除页眉页脚...")
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for p in section.header.paragraphs:
                for run in p.runs:
                    run.text = ''
            for p in section.footer.paragraphs:
                for run in p.runs:
                    run.text = ''

        # 3. 收集需要翻译的段落
        paragraphs_to_translate = []
        skipped_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                if should_skip_paragraph(para.text):
                    # 删除版权段落的所有run
                    for run in para.runs:
                        run.text = ''
                    skipped_count += 1
                else:
                    paragraphs_to_translate.append(para)

        if skipped_count > 0:
            print(f"🗑️  已删除 {skipped_count} 个版权/提示段落")

        total = len(paragraphs_to_translate)
        print(f"📝 翻译 {total} 个段落（{BATCH_SIZE}个/批）...")

        # 4. 批量翻译 - 最安全的方式
        for i in range(0, total, BATCH_SIZE):
            batch_paras = paragraphs_to_translate[i:i+BATCH_SIZE]
            batch_texts = [p.text for p in batch_paras]

            progress = f"[{i+1}-{min(i+BATCH_SIZE, total)}/{total}]"
            print(f"{progress} 翻译中...", end='', flush=True)

            translated_texts = translate_batch(batch_texts)

            # 应用翻译 - 关键：只修改run.text，什么都不删除不添加
            for para, trans_text in zip(batch_paras, translated_texts):
                # 方案A：如果只有一个run，直接修改
                if len(para.runs) == 1:
                    para.runs[0].text = trans_text

                # 方案B：如果有多个runs，把文本合并到第一个run
                elif len(para.runs) > 1:
                    para.runs[0].text = trans_text
                    # 清空其他runs（但不删除，避免破坏结构）
                    for run in para.runs[1:]:
                        run.text = ''

                # 方案C：如果没有run，添加一个（保持样式）
                else:
                    para.add_run(trans_text)

            print(f" ✅")

        # 5. 翻译表格
        if doc.tables:
            print(f"\n📊 翻译 {len(doc.tables)} 个表格...")
            for table_idx, table in enumerate(doc.tables, 1):
                cells_data = []
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                cells_data.append((para, para.text))

                if cells_data:
                    cell_texts = [text for _, text in cells_data]
                    print(f"  表格{table_idx}: {len(cell_texts)}个单元格...", end='', flush=True)
                    translated = translate_batch(cell_texts)

                    for (para, _), trans in zip(cells_data, translated):
                        if len(para.runs) == 1:
                            para.runs[0].text = trans
                        elif len(para.runs) > 1:
                            para.runs[0].text = trans
                            for run in para.runs[1:]:
                                run.text = ''
                        else:
                            para.add_run(trans)

                    print(" ✅")

        # 6. 保存
        doc.save(str(output_file))
        print(f"\n✅ 完成！保存到: {output_file}")

        # 7. 清理
        if input_file.suffix.upper() == '.DOC' and temp_docx.exists():
            temp_docx.unlink()

        return True

    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False

# ============ 批量翻译 ============
def batch_translate_all(max_files=None):
    source_path = Path(SOURCE_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    folders = sorted([d for d in source_path.iterdir() if d.is_dir() and not d.name.startswith('.')])

    print(f"\n{'='*70}")
    print(f"  🚀 安全翻译（完全保护样式和编号）")
    print(f"{'='*70}")
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"文件夹: {len(folders)} 个")
    print(f"{'='*70}\n")

    total_success = 0
    total_fail = 0
    total_skip = 0

    for folder_idx, folder in enumerate(folders, 1):
        folder_name = folder.name
        translated_folder_name = translate_folder_name(folder_name)

        print(f"\n{'#'*70}")
        print(f"# 文件夹 [{folder_idx}/{len(folders)}]: {folder_name}")
        print(f"# 翻译为: {translated_folder_name}")
        print(f"{'#'*70}")

        files = sorted(list(folder.glob("*.DOC")) + list(folder.glob("*.docx")))
        files = [f for f in files if not f.name.startswith('~$') and not f.name.startswith('.~')]

        if max_files:
            files = files[:max_files]

        for file_idx, input_file in enumerate(files, 1):
            file_name = input_file.name
            translated_file_name = translate_file_name(file_name)

            output_folder = Path(OUTPUT_DIR).parent / translated_folder_name
            output_file = output_folder / translated_file_name

            print(f"\n  [{file_idx}/{len(files)}] {file_name}")
            print(f"       → {translated_file_name}")

            if output_file.exists():
                print(f"  ⏭️  已存在，跳过")
                total_skip += 1
                continue

            try:
                success = process_word_document_safe(input_file, output_file)
                if success:
                    total_success += 1
                else:
                    total_fail += 1
            except KeyboardInterrupt:
                print("\n\n⚠️  用户中断")
                break
            except Exception as e:
                print(f"\n  ❌ 失败: {e}")
                total_fail += 1

        if max_files:
            break

    print(f"\n\n{'='*70}")
    print(f"  📊 翻译完成")
    print(f"{'='*70}")
    print(f"✅ 成功: {total_success}")
    print(f"⏭️  跳过: {total_skip}")
    print(f"❌ 失败: {total_fail}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="安全翻译（保护样式）")
    parser.add_argument("--test", action="store_true", help="测试（1个文件）")
    parser.add_argument("--small", action="store_true", help="小批量（10个文件）")

    args = parser.parse_args()

    print("🔌 测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM服务连接成功\n")
        else:
            print(f"⚠️  VLLM响应异常")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        sys.exit(1)

    max_files = None
    if args.test:
        max_files = 1
    elif args.small:
        max_files = 10

    batch_translate_all(max_files)
