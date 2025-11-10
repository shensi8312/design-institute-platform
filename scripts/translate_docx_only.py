#!/usr/bin/env python3
"""
DOCX翻译脚本（完整版）
- 保留所有格式（粗体、斜体、字体等）
- 过滤版权信息
- 删除文本框
- 正确翻译文件名
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
import requests
from datetime import datetime

# 配置
SOURCE_DIR = "specs_docx"
OUTPUT_DIR = "specs_zh"
TRANSLATION_API = "http://10.10.18.3:8000/v1/chat/completions"
MODEL_NAME = "qwen-next-80b"

# 版权关键词（跳过翻译）
COPYRIGHT_KEYWORDS = [
    "Copyright", "©", "AIA", "Deltek",
    "All rights reserved", "版权所有"
]

def translate_text(text, is_filename=False):
    """使用VLLM服务翻译文本"""
    if not text.strip():
        return text

    # 检查是否包含版权信息
    if any(keyword in text for keyword in COPYRIGHT_KEYWORDS):
        return None  # 不翻译版权信息

    try:
        if is_filename:
            prompt = f"请将以下英文直接翻译成中文，只翻译文字本身，不要添加任何额外的词（如'文件'等）：\n\n{text}"
        else:
            prompt = f"请将以下英文翻译成中文，保持专业术语准确性：\n\n{text}"

        response = requests.post(TRANSLATION_API, json={
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 4000,
            "temperature": 0.3
        }, timeout=120)

        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        else:
            print(f"  ❌ API错误: {response.status_code}")
            return None
    except Exception as e:
        print(f"  ❌ 翻译错误: {e}")
        return None

def translate_filename(filename):
    """翻译文件名，保留编号和格式"""
    name_without_ext = filename.rsplit('.', 1)[0]
    ext = '.' + filename.rsplit('.', 1)[1] if '.' in filename else ''

    # 匹配编号部分（如 "071113 FL -"）
    match = re.match(r'^([\d\.\s]+(?:FL|SL|FS)?[\s\-]*)', name_without_ext)

    if match:
        prefix = match.group(1)
        description = name_without_ext[len(prefix):].strip()

        # 翻译描述部分
        if description:
            translated_desc = translate_text(description, is_filename=True)
            if translated_desc:
                return f"{prefix}{translated_desc}{ext}"

    return filename

def remove_textboxes(doc):
    """删除文档中的所有文本框"""
    try:
        # 删除shape中的文本框
        for shape in doc.inline_shapes:
            shape._element.getparent().remove(shape._element)

        # 删除浮动文本框
        for section in doc.sections:
            for element in section._element.xpath('.//w:txbxContent'):
                element.getparent().getparent().getparent().remove(element.getparent().getparent())
    except Exception as e:
        print(f"    ⚠️  文本框删除警告: {e}")

def translate_docx(input_file, output_file):
    """翻译DOCX文件内容（保留格式）"""
    try:
        doc = Document(input_file)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # 删除文本框
        remove_textboxes(doc)

        # 统计需要翻译的runs
        total_runs = sum(len([r for r in para.runs if r.text.strip()]) for para in doc.paragraphs)
        translated = 0
        skipped_copyright = 0

        # 遍历段落和runs（保留格式）
        for para in doc.paragraphs:
            # 检查整个段落是否是版权信息
            if any(keyword in para.text for keyword in COPYRIGHT_KEYWORDS):
                skipped_copyright += len([r for r in para.runs if r.text.strip()])
                para.clear()  # 删除版权段落
                continue

            for run in para.runs:
                if run.text.strip():
                    original_text = run.text
                    translated_text = translate_text(original_text)

                    if translated_text:
                        run.text = translated_text
                        translated += 1
                    else:
                        # 版权信息，删除
                        run.text = ""
                        skipped_copyright += 1

                    if translated % 10 == 0:
                        print(f"      进度: {translated}/{total_runs}", end='\r')

        doc.save(output_file)
        print(f"      进度: {translated}/{total_runs} ✅ (跳过版权: {skipped_copyright})")
        return True
    except Exception as e:
        print(f"    ❌ 处理错误: {e}")
        return False

def main():
    source_path = Path(SOURCE_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    folders = sorted([d for d in source_path.iterdir() if d.is_dir() and not d.name.startswith('.')])

    print(f"\n{'='*70}")
    print(f"  🌐 批量翻译 DOCX → 中文（完整版）")
    print(f"{'='*70}")
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"模型: {MODEL_NAME}")
    print(f"文件夹: {len(folders)} 个")
    print(f"功能: ✅格式保留 ✅过滤版权 ✅删除文本框")
    print(f"{'='*70}\n")

    total_success = 0
    total_fail = 0
    total_skip = 0

    for folder_idx, folder in enumerate(folders, 1):
        folder_name = folder.name

        print(f"\n{'#'*70}")
        print(f"# [{folder_idx}/{len(folders)}] {folder_name}")
        print(f"{'#'*70}")

        files = sorted(list(folder.glob("*.docx")) + list(folder.glob("*.DOCX")))
        files = [f for f in files if not f.name.startswith('~$') and not f.name.startswith('.')]

        print(f"找到 {len(files)} 个DOCX文件\n")

        for file_idx, input_file in enumerate(files, 1):
            file_name = input_file.name

            print(f"  [{file_idx}/{len(files)}] {file_name}")

            # 翻译文件名
            print(f"    📝 翻译文件名...")
            translated_filename = translate_filename(file_name)

            if translated_filename != file_name:
                print(f"    → {translated_filename}")
            else:
                print(f"    → (保持原名)")

            output_folder = Path(OUTPUT_DIR) / folder_name
            output_file = output_folder / translated_filename

            if output_file.exists():
                print(f"    ⏭️  已存在\n")
                total_skip += 1
                continue

            print(f"    🌐 翻译内容...")

            if translate_docx(input_file, output_file):
                total_success += 1
                print()
            else:
                total_fail += 1
                print()

    print(f"\n{'='*70}")
    print(f"  📊 翻译完成")
    print(f"{'='*70}")
    print(f"✅ 成功: {total_success}")
    print(f"⏭️  跳过: {total_skip}")
    print(f"❌ 失败: {total_fail}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    main()
