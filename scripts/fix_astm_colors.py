#!/usr/bin/env python3
"""
修复已翻译文档中ASTM标准编号的颜色问题
移除ASTM编号中M后缀的蓝绿色
"""

import sys
import re
from docx import Document
from docx.shared import RGBColor

def fix_astm_colors(input_path: str, output_path: str):
    """修复ASTM颜色问题"""
    print(f"\n📝 修复文档: {input_path}")
    print(f"输出到: {output_path}\n")

    doc = Document(input_path)
    TEAL_COLOR = RGBColor(0, 176, 176)

    fixed_count = 0

    for para_idx, para in enumerate(doc.paragraphs, 1):
        text = para.text

        # 只处理包含ASTM的段落
        if 'ASTM' not in text:
            continue

        # 检查是否有颜色问题
        has_color_issue = False
        for run in para.runs:
            if run.font.color and run.font.color.rgb == TEAL_COLOR:
                # 检查是否是ASTM相关的run
                if re.search(r'\d+M', run.text):
                    has_color_issue = True
                    break

        if not has_color_issue:
            continue

        # 重建段落，合并所有run的文本
        full_text = para.text

        # 清空段落
        for run in para.runs:
            run.text = ''

        # 找到所有ASTM标准编号（包括完整的格式）
        astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'

        # 分割文本，保留ASTM编号
        parts = []
        last_end = 0

        for match in re.finditer(astm_pattern, full_text):
            # 添加ASTM之前的文本
            if match.start() > last_end:
                parts.append(('text', full_text[last_end:match.start()]))

            # 添加ASTM标准编号
            parts.append(('astm', match.group(0)))
            last_end = match.end()

        # 添加最后剩余的文本
        if last_end < len(full_text):
            parts.append(('text', full_text[last_end:]))

        # 重建runs，ASTM部分不应用颜色
        for part_type, part_text in parts:
            if part_type == 'astm':
                # ASTM标准编号，不应用任何颜色
                run = para.add_run(part_text)
                # 确保没有颜色
                if run.font.color:
                    run.font.color.rgb = None
            else:
                # 普通文本，检查是否有单位需要颜色
                # 分割出单位部分
                unit_pattern = r'\d+(?:\.\d+)?\s*(?:kg/sq\.\s*m|kg/cu\.\s*m|kg/m²|kg/m³|N/mm²|W/m²|mm²|m²|m³|kPa|MPa|mm|cm|km|ml|Pa|°C)(?![²³/a-z])|\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])'

                subparts = re.split(f'({unit_pattern})', part_text, flags=re.IGNORECASE)

                for subpart in subparts:
                    if not subpart:
                        continue

                    run = para.add_run(subpart)

                    # 检查是否是单位
                    if re.match(unit_pattern, subpart, re.IGNORECASE):
                        # 应用蓝绿色
                        run.font.color.rgb = TEAL_COLOR

        fixed_count += 1
        print(f"✅ 修复段落 {para_idx}: {text[:60]}...")

    # 保存文档
    doc.save(output_path)

    print(f"\n🎉 完成！共修复 {fixed_count} 个段落")
    print(f"输出文件: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("使用方法: python3 fix_astm_colors.py <输入文件> <输出文件>")
        sys.exit(1)

    fix_astm_colors(sys.argv[1], sys.argv[2])
