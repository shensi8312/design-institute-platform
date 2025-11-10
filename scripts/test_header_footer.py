#!/usr/bin/env python3
"""
测试Word文档的页眉页脚删除功能
"""

import sys
from pathlib import Path
from docx import Document

def check_header_footer(docx_file: Path):
    """检查Word文档的页眉页脚"""
    print(f"\n{'='*70}")
    print(f"📄 检查文件: {docx_file.name}")
    print(f"{'='*70}")

    try:
        doc = Document(str(docx_file))

        total_sections = len(doc.sections)
        print(f"\n📊 文档共有 {total_sections} 个节（Section）\n")

        has_header = False
        has_footer = False

        for i, section in enumerate(doc.sections, 1):
            print(f"第 {i} 节:")

            # 检查页眉
            header_paras = section.header.paragraphs
            header_text = "\n".join([p.text for p in header_paras if p.text.strip()])

            if header_paras and any(p.text.strip() for p in header_paras):
                print(f"  ❌ 页眉存在（{len(header_paras)} 个段落）")
                print(f"     内容: {header_text[:100]}...")
                has_header = True
            else:
                print(f"  ✅ 页眉已删除")

            # 检查页脚
            footer_paras = section.footer.paragraphs
            footer_text = "\n".join([p.text for p in footer_paras if p.text.strip()])

            if footer_paras and any(p.text.strip() for p in footer_paras):
                print(f"  ❌ 页脚存在（{len(footer_paras)} 个段落）")
                print(f"     内容: {footer_text[:100]}...")
                has_footer = True
            else:
                print(f"  ✅ 页脚已删除")

            print()

        # 总结
        print(f"{'='*70}")
        if not has_header and not has_footer:
            print(f"✅ 页眉页脚删除成功！")
        else:
            if has_header:
                print(f"❌ 仍有页眉存在")
            if has_footer:
                print(f"❌ 仍有页脚存在")
        print(f"{'='*70}\n")

        return not has_header and not has_footer

    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 test_header_footer.py <docx文件路径>")
        print("\n示例:")
        print("  python3 test_header_footer.py 'docs/specs_zh/全长规范/000101 FL - 项目标题页.docx'")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)

    if not docx_path.suffix.lower() in ['.docx', '.doc']:
        print(f"❌ 不是Word文档: {docx_path}")
        sys.exit(1)

    success = check_header_footer(docx_path)
    sys.exit(0 if success else 1)
