#!/usr/bin/env python3
"""
调试Word文档的列表编号结构
"""

import sys
from pathlib import Path
from docx import Document

def debug_list_structure(docx_file: Path):
    """检查Word文档的列表结构"""
    print(f"\n{'='*70}")
    print(f"📄 检查文件: {docx_file.name}")
    print(f"{'='*70}\n")

    try:
        doc = Document(str(docx_file))

        # 查找有列表编号的段落
        list_paras = []
        for i, para in enumerate(doc.paragraphs, 1):
            # 检查段落样式
            if para.style and para.style.name:
                style_name = para.style.name

                # 检查是否是列表样式
                if any(keyword in style_name.lower() for keyword in ['list', 'number', 'bullet']):
                    list_paras.append({
                        'line': i,
                        'style': style_name,
                        'text': para.text[:100],
                        'format': para.paragraph_format
                    })

            # 或者文本开头有编号
            text = para.text.strip()
            if text and (text[0].isdigit() or (len(text) > 1 and text[1] == '.')):
                if para not in [p for p in list_paras]:
                    list_paras.append({
                        'line': i,
                        'style': para.style.name if para.style else 'None',
                        'text': text[:100],
                        'format': para.paragraph_format
                    })

        print(f"找到 {len(list_paras)} 个列表段落\n")

        for item in list_paras[:15]:  # 显示前15个
            print(f"[{item['line']}] 样式: {item['style']}")
            print(f"     文本: {item['text']}")
            print(f"     格式: 左缩进={item['format'].left_indent}, 首行缩进={item['format'].first_line_indent}")
            print()

        # 检查一个具体段落的XML结构
        if list_paras:
            test_para = doc.paragraphs[list_paras[0]['line'] - 1]
            print(f"{'='*70}")
            print(f"示例段落的XML结构（前500字符）:")
            print(f"{'='*70}")
            xml_str = test_para._element.xml.decode('utf-8') if isinstance(test_para._element.xml, bytes) else str(test_para._element.xml)
            print(xml_str[:500])
            print("...")

        return True

    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 debug_list_numbers.py <docx文件路径>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)

    debug_list_structure(docx_path)
