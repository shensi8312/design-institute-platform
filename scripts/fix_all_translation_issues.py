#!/usr/bin/env python3
"""
一次性修复所有翻译问题
"""

import sys
import re
from docx import Document
from docx.shared import RGBColor

def fix_all_issues(input_path: str, output_path: str):
    """修复所有翻译问题"""
    print(f"\n📝 修复所有翻译问题: {input_path}")
    print(f"输出到: {output_path}\n")

    doc = Document(input_path)
    TEAL_COLOR = RGBColor(0, 176, 176)

    # 1. 术语替换字典
    term_replacements = {
        '屋顶重新覆盖': '屋顶翻新',
        '重新覆盖': '翻新',
        '重新铺设': '翻新',
        '重新做屋面': '屋顶翻新',
        '重新屋顶': '屋顶翻新',
        '进行重新屋顶施工': '进行屋顶翻新施工',
        '待重新屋顶施工': '待翻新',
        '重新屋顶施工': '屋顶翻新施工',
        '不进行重新铺设': '不进行翻新',
        '需重新铺设屋顶': '需翻新屋顶',
    }

    # 2. 占位符翻译
    placeholder_translations = {
        'location': '地点',
        'dimension': '尺寸',
        'area': '面积',
        'value': '数值',
        'method': '方法',
        'name': '名称',
    }

    # 3. 章节标题翻译
    section_translations = {
        'Coated Foamed Roofing': '涂覆泡沫屋面',
        'Sheet Metal Flashing and Trim': '金属泛水和饰边',
        'Roof Specialties': '屋面专项',
    }

    fixed_count = 0

    for para_idx, para in enumerate(doc.paragraphs, 1):
        original_text = para.text
        modified = False

        # 检查是否包含ASTM且有颜色问题
        has_astm_color_issue = False
        if 'ASTM' in original_text:
            for run in para.runs:
                if run.font.color and run.font.color.rgb == TEAL_COLOR:
                    if re.search(r'\d+M', run.text):
                        has_astm_color_issue = True
                        break

        # 如果有ASTM颜色问题，重建段落
        if has_astm_color_issue:
            full_text = para.text
            for run in para.runs:
                run.text = ''

            astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'
            parts = []
            last_end = 0

            for match in re.finditer(astm_pattern, full_text):
                if match.start() > last_end:
                    parts.append(('text', full_text[last_end:match.start()]))
                parts.append(('astm', match.group(0)))
                last_end = match.end()

            if last_end < len(full_text):
                parts.append(('text', full_text[last_end:]))

            for part_type, part_text in parts:
                if part_type == 'astm':
                    run = para.add_run(part_text)
                    if run.font.color:
                        run.font.color.rgb = None
                else:
                    run = para.add_run(part_text)

            modified = True

        # 获取当前文本（可能已被ASTM修复修改）
        current_text = para.text

        # 应用术语替换
        new_text = current_text
        for old, new in term_replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
                modified = True

        # 翻译占位符内容
        def translate_placeholder(match):
            content = match.group(1)
            for en, zh in placeholder_translations.items():
                if content.lower() == en.lower():
                    return f'〈待填：{zh}〉'
            return match.group(0)

        placeholder_pattern = r'〈待填：([^〉]+)〉'
        translated_text = re.sub(placeholder_pattern, translate_placeholder, new_text)
        if translated_text != new_text:
            new_text = translated_text
            modified = True

        # 翻译章节标题
        for en, zh in section_translations.items():
            if en in new_text:
                new_text = new_text.replace(en, zh)
                modified = True

        # 如果文本被修改（除了ASTM颜色修复），重建段落
        if modified and new_text != current_text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

        if modified:
            fixed_count += 1
            print(f"✅ 修正段落 {para_idx}")

    # 保存文档
    doc.save(output_path)

    print(f"\n🎉 完成！共修正 {fixed_count} 个段落")
    print(f"输出文件: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("使用方法: python3 fix_all_translation_issues.py <输入文件> <输出文件>")
        sys.exit(1)

    fix_all_issues(sys.argv[1], sys.argv[2])
