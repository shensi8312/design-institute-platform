#!/usr/bin/env python3
"""
统一"reroofing"相关术语的翻译
将"重新"改为"翻新"
"""

import sys
from docx import Document

def fix_reroofing_terms(input_path: str, output_path: str):
    """修正屋顶翻新术语"""
    print(f"\n📝 修正文档: {input_path}")
    print(f"输出到: {output_path}\n")

    doc = Document(input_path)

    # 定义替换规则
    replacements = {
        '屋顶重新覆盖': '屋顶翻新',
        '重新覆盖': '翻新',
        '重新铺设': '翻新',
        '重新做屋面': '屋顶翻新',
        '重新屋顶': '屋顶翻新',
        '进行重新屋顶施工': '进行屋顶翻新施工',
        '待重新屋顶施工': '待翻新',
        '重新屋顶施工': '屋顶翻新施工',
        '不进行重新铺设': '不进行翻新',
        '需重新铺设屋顶': '需翻新屋顶',
    }

    fixed_count = 0

    for para_idx, para in enumerate(doc.paragraphs, 1):
        original_text = para.text

        # 检查是否需要替换
        needs_fix = any(old in original_text for old in replacements.keys())

        if not needs_fix:
            continue

        # 进行替换
        new_text = original_text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)

        if new_text != original_text:
            # 重建段落保持格式
            # 清空现有runs
            for run in para.runs:
                run.text = ''

            # 添加新文本（使用第一个run保持格式）
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

            fixed_count += 1
            print(f"✅ 修正段落 {para_idx}:")
            print(f"   原: {original_text[:60]}...")
            print(f"   新: {new_text[:60]}...")
            print()

    # 保存文档
    doc.save(output_path)

    print(f"\n🎉 完成！共修正 {fixed_count} 个段落")
    print(f"输出文件: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("使用方法: python3 fix_reroofing_translation.py <输入文件> <输出文件>")
        sys.exit(1)

    fix_reroofing_terms(sys.argv[1], sys.argv[2])
