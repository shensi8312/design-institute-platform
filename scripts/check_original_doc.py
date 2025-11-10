#!/usr/bin/env python3
"""
检查原始DOC文档的列表编号结构
"""

import sys
from pathlib import Path
from docx import Document

def check_original(docx_file: Path):
    """检查原始文档"""
    print(f"\n{'='*70}")
    print(f"📄 检查原始文档: {docx_file.name}")
    print(f"{'='*70}\n")

    try:
        doc = Document(str(docx_file))

        # 查找第41-50段（问题区域）
        for i in range(40, min(51, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if para.text.strip():
                print(f"[{i+1}] {para.text[:120]}")
                print(f"     样式: {para.style.name if para.style else 'None'}")
                print()

        return True

    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 check_original_doc.py <docx文件路径>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)

    check_original(docx_path)
