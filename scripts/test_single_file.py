#!/usr/bin/env python3
"""
测试单个文件的转换和翻译
"""

import os
import sys
import subprocess
from pathlib import Path

# 添加父目录到路径以导入batch_convert_and_translate
sys.path.insert(0, str(Path(__file__).parent))

from batch_convert_and_translate import (
    convert_doc_to_docx,
    translate_docx,
    translate_file_name,
    VLLM_URL
)
import requests

# 测试文件
TEST_FILE = "docs/specs/Full Length/00 - PROCUREMENT AND CONTRACTING REQUIREMENTS/000107 FL - Seals Page.DOC"
OUTPUT_DIR = "docs/specs_zh/00 - 采购与合同要求"

def test_single_file():
    print("\n" + "="*70)
    print("  🧪 单文件测试")
    print("="*70)

    input_file = Path(TEST_FILE)
    if not input_file.exists():
        print(f"❌ 文件不存在: {TEST_FILE}")
        return

    # 翻译文件名
    print(f"\n📝 原文件名: {input_file.name}")
    translated_name = translate_file_name(input_file.name)
    print(f"   翻译后: {translated_name}")

    # 输出文件
    output_folder = Path(OUTPUT_DIR)
    output_file = output_folder / translated_name

    if output_file.exists():
        print(f"\n🗑️  删除旧文件: {output_file}")
        output_file.unlink()

    # 步骤1：转换
    print(f"\n🔄 步骤1: DOC → DOCX 转换...")
    if convert_doc_to_docx(input_file, output_file):
        print(f"   ✅ 转换成功")
    else:
        print(f"   ❌ 转换失败")
        return

    # 步骤2：翻译
    print(f"\n📝 步骤2: 翻译内容...")
    if translate_docx(output_file):
        print(f"   ✅ 翻译成功")
    else:
        print(f"   ❌ 翻译失败")
        return

    # 步骤3：验证
    print(f"\n🔍 步骤3: 验证结果...")
    verify_result(output_file)

    print(f"\n✅ 测试完成！")
    print(f"   输出文件: {output_file}")

def verify_result(file_path):
    """验证翻译结果"""
    from docx import Document

    doc = Document(str(file_path))

    # 统计段落
    total_paras = len(doc.paragraphs)
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]

    # 统计样式
    styles = {}
    for para in doc.paragraphs:
        style = para.style.name
        if style not in styles:
            styles[style] = 0
        styles[style] += 1

    # 检查空段落
    empty_cmt_tip = [p for p in doc.paragraphs if not p.text.strip() and p.style.name in ["CMT", "TIP"]]

    # 检查是否有英文残留
    english_paras = []
    for para in non_empty_paras:
        text = para.text.lower()
        if any(keyword in text for keyword in ["copyright", "double click", "tips:", "aia"]):
            english_paras.append(para.text[:80])

    print(f"   总段落数: {total_paras}")
    print(f"   非空段落: {len(non_empty_paras)}")
    print(f"   样式分布: {styles}")

    if empty_cmt_tip:
        print(f"   ⚠️  空的CMT/TIP段落: {len(empty_cmt_tip)}个")
    else:
        print(f"   ✅ 无空段落")

    if english_paras:
        print(f"   ⚠️  英文残留:")
        for text in english_paras[:3]:
            print(f"      - {text}")
    else:
        print(f"   ✅ 无英文残留")

if __name__ == "__main__":
    # 测试VLLM连接
    print("🔌 测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM连接成功\n")
        else:
            print(f"⚠️  VLLM响应异常")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        sys.exit(1)

    test_single_file()
