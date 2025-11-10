#!/usr/bin/env python3
"""
验证翻译结果 - 检查所有已知问题是否修复
"""

import sys
from pathlib import Path
from docx import Document
import re

def validate_translation(docx_path: str):
    """验证翻译质量"""

    print("\n" + "="*70)
    print("📋 翻译验证报告")
    print("="*70)
    print(f"文件: {docx_path}\n")

    doc = Document(docx_path)

    issues = []
    warnings = []
    successes = []

    # 1. 检查列表编号问题
    print("1️⃣ 检查列表编号...")
    numbering_issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 检查是否有 "a.2." "b.3." 这种格式
        if re.search(r'^[a-z]\.\d+\.', text):
            numbering_issues.append(f"  段落 {i+1}: {text[:80]}")

        # 检查是否有纯数字编号开头（VLLM添加的）
        if re.match(r'^\d+[\.\)、]\s+', text) and para.style and any(s in para.style.name.upper() for s in ['PR', 'LIST', 'CMT']):
            numbering_issues.append(f"  段落 {i+1}: {text[:80]}")

    if numbering_issues:
        issues.append("❌ 发现列表编号问题:")
        issues.extend(numbering_issues[:5])  # 只显示前5个
        if len(numbering_issues) > 5:
            issues.append(f"  ... 还有 {len(numbering_issues)-5} 个问题")
    else:
        successes.append("✅ 列表编号正确：无 a.2. b.3. 格式")

    # 2. 检查章节标题翻译
    print("2️⃣ 检查章节标题翻译...")
    section_checks = {
        "ACTION SUBMITTALS": "操作提交件",
        "INFORMATIONAL SUBMITTALS": "信息提交件",
        "CLOSEOUT SUBMITTALS": "竣工提交件",
    }

    found_sections = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        for eng, chn in section_checks.items():
            if eng in text.upper():
                found_sections[eng] = ("英文未翻译", text)
            elif chn in text:
                found_sections[eng] = ("已翻译", text)

    for eng, chn in section_checks.items():
        if eng in found_sections:
            status, text = found_sections[eng]
            if status == "英文未翻译":
                issues.append(f"❌ '{eng}' 未翻译: {text[:50]}")
            else:
                successes.append(f"✅ '{eng}' → '{chn}'")

    # 3. 检查 END OF SECTION
    print("3️⃣ 检查 END OF SECTION 翻译...")
    end_of_section_found = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if "END OF SECTION" in text.upper() and not "节结束" in text:
            issues.append(f"❌ END OF SECTION 未翻译: {text}")
            end_of_section_found = True
        elif "节结束" in text:
            successes.append(f"✅ END OF SECTION 已翻译: {text}")
            end_of_section_found = True

    if not end_of_section_found:
        warnings.append("⚠️ 未找到 END OF SECTION 文本")

    # 4. 检查页脚
    print("4️⃣ 检查页脚...")
    footer_texts = []
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for para in footer.paragraphs:
                if para.text.strip():
                    footer_texts.append(para.text.strip())

    if footer_texts:
        issues.append(f"❌ 发现页脚内容: {footer_texts}")
    else:
        successes.append("✅ 所有页脚已清空")

    # 5. 检查单位格式和颜色
    print("5️⃣ 检查单位格式...")
    unit_issues = []
    unit_count = 0
    for para in doc.paragraphs:
        # 检查是否有单位
        if re.search(r'\d+\s*(kg/m²|m²|mm|℃|MPa|kPa)', para.text):
            unit_count += 1
            # 检查颜色
            has_color = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    if r == 0 and g == 176 and b == 176:
                        has_color = True
                        break

            if not has_color and 'kg/m²' in para.text:
                unit_issues.append(f"  段落缺少颜色: {para.text[:60]}")

    if unit_count > 0:
        successes.append(f"✅ 找到 {unit_count} 个单位格式")
        if unit_issues:
            warnings.extend(["⚠️ 部分单位缺少颜色标记:"] + unit_issues[:3])

    # 6. 检查占位符残留
    print("6️⃣ 检查占位符残留...")
    placeholder_issues = []
    for i, para in enumerate(doc.paragraphs):
        if '⟨⟨UNIT' in para.text or 'MARKER' in para.text or '<<UNIT' in para.text:
            placeholder_issues.append(f"  段落 {i+1}: {para.text[:80]}")

    if placeholder_issues:
        issues.append("❌ 发现占位符残留:")
        issues.extend(placeholder_issues[:5])
    else:
        successes.append("✅ 无占位符残留")

    # 打印报告
    print("\n" + "="*70)
    print("📊 验证结果")
    print("="*70 + "\n")

    if successes:
        print("✅ 成功项:")
        for s in successes:
            print(f"  {s}")
        print()

    if warnings:
        print("⚠️ 警告项:")
        for w in warnings:
            print(f"  {w}")
        print()

    if issues:
        print("❌ 问题项:")
        for issue in issues:
            print(f"  {issue}")
        print()
        print(f"总计: {len(issues)} 个问题")
        return False
    else:
        print("🎉 所有检查通过！翻译质量良好。")
        return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 validate_translation.py <docx文件>")
        sys.exit(1)

    docx_path = sys.argv[1]
    success = validate_translation(docx_path)
    sys.exit(0 if success else 1)
