#!/usr/bin/env python3
"""
修复占位符内部的英文翻译
"""

import sys
import re
from docx import Document

def fix_placeholder_content(input_path: str, output_path: str):
    """修复占位符内容"""
    print(f"\n📝 修正占位符: {input_path}")
    print(f"输出到: {output_path}\n")

    doc = Document(input_path)

    # 占位符翻译词典
    placeholder_translations = {
        'location': '地点',
        'dimension': '尺寸',
        'dimensions': '尺寸',
        'area': '面积',
        'value': '数值',
        'values': '数值',
        'method': '方法',
        'name': '名称',
        'type': '类型',
        'thickness': '厚度',
        'width': '宽度',
        'height': '高度',
        'length': '长度',
        'depth': '深度',
        'spacing': '间距',
        'size': '尺寸',
        'quantity': '数量',
        'color': '颜色',
        'material': '材料',
        'finish': '完成面',
        'pattern': '图案',
    }

    fixed_count = 0

    for para_idx, para in enumerate(doc.paragraphs, 1):
        original_text = para.text

        # 查找占位符 〈待填：xxx〉
        pattern = r'〈待填：([^〉]+)〉'

        if not re.search(pattern, original_text):
            continue

        new_text = original_text

        def translate_placeholder(match):
            content = match.group(1)
            # 翻译占位符内容
            translated = content
            for en, zh in placeholder_translations.items():
                # 只翻译纯英文单词
                if content.lower() == en.lower():
                    translated = zh
                    break
            return f'〈待填：{translated}〉'

        new_text = re.sub(pattern, translate_placeholder, new_text)

        if new_text != original_text:
            # 重建段落
            for run in para.runs:
                run.text = ''

            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

            fixed_count += 1
            print(f"✅ 修正段落 {para_idx}:")
            print(f"   原: {original_text[:80]}...")
            print(f"   新: {new_text[:80]}...")
            print()

    # 保存文档
    doc.save(output_path)

    print(f"\n🎉 完成！共修正 {fixed_count} 个段落")
    print(f"输出文件: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("使用方法: python3 fix_placeholder_translations.py <输入文件> <输出文件>")
        sys.exit(1)

    fix_placeholder_content(sys.argv[1], sys.argv[2])
