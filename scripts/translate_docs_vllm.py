#!/usr/bin/env python3
"""
批量文档翻译脚本 - 使用VLLM + Qwen3-32B
不需要API Token，直接调用本地VLLM服务
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import List, Dict, Optional
import requests

# 强制刷新输出，确保实时显示
sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# 配置
SOURCE_DIR = "docs/specs"
OUTPUT_DIR = "docs/specs_zh"
# 使用本地 VLLM 服务 (Qwen3-32B)
VLLM_URL = os.getenv("VLLM_URL", "http://10.10.18.3:8000")
VLLM_MODEL = os.getenv("VLLM_MODEL", "/mnt/data/models/Qwen3-32B")

def translate_text_vllm(text: str, target_lang: str = "zh") -> str:
    """
    使用VLLM的Qwen3-32B模型翻译文本
    """
    if not text or not text.strip():
        return text

    # 构建建筑行业专业翻译提示词
    prompt = f"""你是资深的建筑工程规范文档翻译专家，专门翻译AIA MasterSpec等建筑设计行业文档。

翻译要求：
1. **专业术语**：准确翻译建筑专业术语（如ASTM标准、材料名称、施工工艺等）
2. **保留标记**：保持[xxx]、<xxx>等占位符不翻译
3. **单位换算**：保留英制和公制单位（如inches/mm），不做转换
4. **格式完整**：保持原文格式、编号和结构
5. **简洁准确**：只返回翻译结果，不添加说明

请翻译以下建筑规范文本：

{text}

中文翻译："""

    try:
        # 调用VLLM服务
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,  # 降低温度以获得更稳定的翻译
                "max_tokens": 2000
            },
            timeout=120
        )

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            # 移除可能的"翻译："前缀
            if translated.startswith("翻译："):
                translated = translated[3:].strip()

            return translated
        else:
            print(f"VLLM返回错误: {response.status_code} - {response.text}")
            return text

    except requests.exceptions.Timeout:
        print(f"翻译超时，返回原文")
        return text
    except Exception as e:
        print(f"翻译异常: {e}")
        return text

def process_doc_file(input_file: Path, output_file: Path):
    """
    处理DOC/DOCX文件，保留格式，翻译文本，移除页眉页脚
    """
    print(f"\n{'='*60}")
    print(f"处理文档: {input_file.name}")
    print(f"{'='*60}")

    # 第一步：将DOC转换为DOCX（如果是DOC格式）
    if input_file.suffix.upper() == '.DOC':
        temp_docx = Path(f"/tmp/{input_file.stem}_temp.docx")
        # 使用textutil (macOS) 转换
        if sys.platform == 'darwin':
            result = subprocess.run([
                'textutil', '-convert', 'docx',
                str(input_file), '-output', str(temp_docx)
            ], capture_output=True)

            if result.returncode == 0 and temp_docx.exists():
                working_file = temp_docx
            else:
                print(f"警告: 无法转换 {input_file.name}，跳过")
                return
        else:
            print(f"警告: 非macOS系统，无法转换DOC格式，跳过")
            return
    else:
        working_file = input_file

    # 确保输出目录存在
    output_file.parent.mkdir(parents=True, exist_ok=True)

    # 第二步：使用python-docx处理DOCX文件
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document(str(working_file))

        # 移除页眉页脚
        print("移除页眉页脚...")
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # 清空页眉
            for paragraph in section.header.paragraphs:
                paragraph.clear()

            # 清空页脚
            for paragraph in section.footer.paragraphs:
                paragraph.clear()

        # 翻译正文段落，保留格式
        total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        translated_count = 0

        print(f"开始翻译 {total_paragraphs} 个段落...\n")

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                translated_count += 1
                original_text = paragraph.text

                # 显示进度
                print(f"[{translated_count}/{total_paragraphs}] 翻译: {original_text[:60]}...")

                # 保存原有格式
                original_runs = []
                for run in paragraph.runs:
                    original_runs.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name if run.font.name else None,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })

                # 翻译文本
                translated_text = translate_text_vllm(original_text)

                # 清空段落
                paragraph.clear()

                # 重新添加翻译后的文本，应用原有格式
                if original_runs:
                    # 使用第一个run的格式
                    first_format = original_runs[0]
                    run = paragraph.add_run(translated_text)

                    if first_format['bold'] is not None:
                        run.bold = first_format['bold']
                    if first_format['italic'] is not None:
                        run.italic = first_format['italic']
                    if first_format['underline'] is not None:
                        run.underline = first_format['underline']
                    if first_format['font_name']:
                        run.font.name = first_format['font_name']
                    if first_format['font_size']:
                        run.font.size = first_format['font_size']
                    if first_format['font_color']:
                        run.font.color.rgb = first_format['font_color']
                else:
                    paragraph.add_run(translated_text)

        # 翻译表格
        table_count = len(doc.tables)
        if table_count > 0:
            print(f"\n翻译 {table_count} 个表格...")
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"  表格 {table_idx}/{table_count}")
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                translated = translate_text_vllm(original_text)

                                # 保留格式，更新文本
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    bold = first_run.bold
                                    italic = first_run.italic

                                    paragraph.clear()
                                    new_run = paragraph.add_run(translated)
                                    new_run.bold = bold
                                    new_run.italic = italic

        # 保存翻译后的文档
        doc.save(str(output_file))
        print(f"\n✅ 翻译完成！")
        print(f"已保存: {output_file}")

        # 清理临时文件
        if input_file.suffix.upper() == '.DOC' and temp_docx.exists():
            temp_docx.unlink()

    except ImportError:
        print("\n❌ 错误: 需要安装 python-docx 库")
        print("请运行: pip3 install python-docx --break-system-packages")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ 处理文档时出错: {e}")
        import traceback
        traceback.print_exc()

def batch_translate(source_dir: str = SOURCE_DIR, output_dir: str = OUTPUT_DIR, max_files: int = None):
    """
    批量翻译目录下的所有文档
    """
    source_path = Path(source_dir)
    output_path = Path(output_dir)

    if not source_path.exists():
        print(f"❌ 错误: 源目录不存在: {source_dir}")
        return

    # 获取所有DOC/DOCX文件
    doc_files = list(source_path.rglob("*.DOC")) + list(source_path.rglob("*.docx"))

    # 过滤掉临时文件
    doc_files = [f for f in doc_files if not f.name.startswith('~$')]

    total_files = len(doc_files)

    if max_files:
        doc_files = doc_files[:max_files]
        print(f"📋 测试模式：只处理前 {max_files} 个文件")

    print(f"\n{'='*60}")
    print(f"  批量文档翻译")
    print(f"{'='*60}")
    print(f"源目录: {source_dir}")
    print(f"输出目录: {output_dir}")
    print(f"VLLM服务: {VLLM_URL}")
    print(f"模型: {VLLM_MODEL}")
    print(f"找到 {total_files} 个DOC/DOCX文件")
    if max_files:
        print(f"将处理: {len(doc_files)} 个文件（测试模式）")
    else:
        print(f"将处理: {len(doc_files)} 个文件")
    print(f"{'='*60}\n")

    # 处理文件
    success_count = 0
    fail_count = 0

    for i, input_file in enumerate(doc_files, 1):
        print(f"\n\n{'#'*60}")
        print(f"# 进度: [{i}/{len(doc_files)}]")
        print(f"{'#'*60}")

        # 计算相对路径
        rel_path = input_file.relative_to(source_path)
        output_file = output_path / rel_path

        # 确保输出文件是.docx格式
        if output_file.suffix.upper() == '.DOC':
            output_file = output_file.with_suffix('.docx')

        # 跳过已存在的文件
        if output_file.exists():
            print(f"⏭️  已存在，跳过: {output_file.name}")
            continue

        try:
            process_doc_file(input_file, output_file)
            success_count += 1
        except KeyboardInterrupt:
            print("\n\n⚠️  用户中断")
            break
        except Exception as e:
            print(f"\n❌ 失败: {e}")
            fail_count += 1
            continue

    # 统计
    print(f"\n\n{'='*60}")
    print(f"  翻译完成统计")
    print(f"{'='*60}")
    print(f"✅ 成功: {success_count} 个文件")
    print(f"❌ 失败: {fail_count} 个文件")
    print(f"⏭️  跳过: {len(doc_files) - success_count - fail_count} 个文件")
    print(f"输出目录: {output_dir}")
    print(f"{'='*60}\n")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="批量翻译文档 - 使用VLLM+Qwen3-32B")
    parser.add_argument("--source", default=SOURCE_DIR, help="源目录")
    parser.add_argument("--output", default=OUTPUT_DIR, help="输出目录")
    parser.add_argument("--test", action="store_true", help="测试模式（只处理前1个文件）")
    parser.add_argument("--small", action="store_true", help="小批量测试（处理前10个文件）")
    parser.add_argument("--vllm-url", default=VLLM_URL, help="VLLM服务地址")

    args = parser.parse_args()

    # 更新配置
    VLLM_URL = args.vllm_url

    # 测试VLLM连接
    print("测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM服务连接成功: {VLLM_URL}")
        else:
            print(f"⚠️  VLLM服务响应异常: {response.status_code}")
    except Exception as e:
        print(f"❌ VLLM服务连接失败: {e}")
        print("请检查VLLM服务是否正在运行")
        sys.exit(1)

    # 确定处理文件数量
    max_files = None
    if args.test:
        max_files = 1
    elif args.small:
        max_files = 10

    batch_translate(args.source, args.output, max_files)
