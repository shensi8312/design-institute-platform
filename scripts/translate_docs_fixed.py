#!/usr/bin/env python3
"""
完整批量处理脚本
步骤1：用 LibreOffice 转换 DOC → DOCX（保留所有样式）
步骤2：翻译 DOCX（保留所有格式）
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import List, Dict
import requests

sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# ============ 配置 ============
SOURCE_DIR = "docs/specs/Full Length"
OUTPUT_DIR = "docs/specs_zh/Full Length"
TRANSLATION_MAP = "docs/specs/文件名翻译对照表_完整版.json"
VLLM_URL = os.getenv("VLLM_URL", "http://10.10.18.3:8000")
VLLM_MODEL = os.getenv("VLLM_MODEL", "/mnt/data/models/Qwen3-32B")
BATCH_SIZE = 10
LIBREOFFICE = os.getenv("LIBREOFFICE", "soffice")  # Linux: soffice, macOS: /Applications/LibreOffice.app/Contents/MacOS/soffice

# ============ 加载翻译对照表 ============
def load_translation_map() -> Dict:
    try:
        with open(TRANSLATION_MAP, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {"folders": {}, "files": {}}

TRANS_MAP = load_translation_map()

def translate_folder_name(folder_name: str) -> str:
    return TRANS_MAP.get("folders", {}).get(folder_name, folder_name)

def is_translation_complete(text: str) -> bool:
    """检查翻译是否完整（不包含常见英文单词）"""
    # 常见的需要翻译的英文单词
    common_words = [
        'for', 'and', 'of', 'the', 'in', 'to', 'with', 'on', 'at',
        'Preparation', 'Installation', 'System', 'Equipment', 'Materials',
        'Requirements', 'Specifications', 'Design', 'Construction', 'Details',
        'General', 'Special', 'Testing', 'Quality', 'Control', 'Management',
        'Building', 'Structure', 'Mechanical', 'Electrical', 'Plumbing'
    ]

    # 检查是否包含英文单词（忽略数字和"FL"之类的缩写）
    words = text.split()
    for word in words:
        # 跳过数字和短缩写
        if word.isdigit() or len(word) <= 2:
            continue
        # 检查是否是英文单词
        for common in common_words:
            if common.lower() in word.lower():
                return False
    return True

def translate_file_name(file_name: str) -> str:
    """翻译文件名（优先使用对照表，否则用VLLM翻译）"""
    name_without_ext = file_name.replace('.DOC', '').replace('.docx', '')

    # 尝试从对照表获取
    if file_name in TRANS_MAP.get("files", {}):
        translated = TRANS_MAP["files"][file_name]
        # 检查翻译是否完整（不包含常见英文单词）
        if translated != name_without_ext and is_translation_complete(translated):
            return f"{translated}.docx"

    # 如果对照表中没有或翻译不完整，用VLLM翻译
    # 分离编号和名称（如 "000107 FL - Seals Page" -> "000107 FL - " + "Seals Page"）
    parts = name_without_ext.split(' - ', 1)
    if len(parts) == 2:
        prefix = parts[0]  # "000107 FL"
        name_part = parts[1]  # "Seals Page"

        # 只翻译名称部分
        translated_name = translate_single(name_part)
        return f"{prefix} - {translated_name}.docx"

    # 如果格式不符合预期，保持原样
    return f"{name_without_ext}.docx"

# ============ 章节标题翻译字典 ============
SECTION_HEADINGS = {
    "ACTION SUBMITTALS": "操作提交件",
    "INFORMATIONAL SUBMITTALS": "信息提交件",
    "CLOSEOUT SUBMITTALS": "竣工提交件",
    "QUALITY ASSURANCE": "质量保证",
    "DELIVERY, STORAGE, AND HANDLING": "交付、存储和处理",
    "FIELD CONDITIONS": "现场条件",
    "WARRANTY": "保修",
    "MAINTENANCE": "维护",
}

# ============ 内容过滤 ============
def should_skip_paragraph(text: str) -> bool:
    """判断段落是否应该被完全删除"""
    skip_keywords = [
        "Copyright",
        "The American Institute of Architects",
        "AIA",
        "Exclusively published and distributed by Deltek",
        "TIPS:",
        "TIP:",
        "To view non-printing Editor's Notes",
        "MasterWorks/Single-File Formatting",
        "MasterWorks/Supporting Information",
        "Content Requests:",
        "<Double click here",
        "Double click here",
        "submit questions",
        "comments",
        "suggested edits",
    ]
    for keyword in skip_keywords:
        if keyword.lower() in text.lower():
            return True

    # 不再删除SECTION行，改为在预处理时去掉前缀
    return False

def preprocess_text(text: str, style_name: str = None) -> tuple[str, bool]:
    """预处理文本（去掉SECTION前缀，去掉列表编号，删除英制单位，翻译章节标题）

    返回: (处理后的文本, 是否已翻译为中文)
    """
    import re

    # 检查是否是章节标题，直接替换为中文
    text_upper = text.strip().upper()
    if text_upper in SECTION_HEADINGS:
        return (SECTION_HEADINGS[text_upper], True)  # 已翻译

    # 处理 "PART 数字 - 标题" 模式
    # "PART 1 - GENERAL" → "第1部分 - 总则"
    # "PART 2 - PRODUCTS" → "第2部分 - 产品"
    # "PART 3 - EXECUTION" → "第3部分 - 施工"
    part_match = re.match(r'^PART\s+(\d+)\s*-?\s*(.*)$', text_upper)
    if part_match:
        part_num = part_match.group(1)
        part_title = part_match.group(2).strip()
        # 预定义的PART标题
        part_titles = {
            "GENERAL": "总则",
            "PRODUCTS": "产品",
            "EXECUTION": "施工",
            "SUBMITTALS": "提交件",
        }
        if part_title in part_titles:
            return (f"第{part_num}部分 - {part_titles[part_title]}", True)
        elif part_title:
            # 有标题但不在预定义中，翻译标题
            translated_title = translate_single(part_title)
            return (f"第{part_num}部分 - {translated_title}", True)
        else:
            # 只有PART数字
            return (f"第{part_num}部分", True)

    # 处理 "END OF SECTION 数字" 模式
    end_of_section_match = re.match(r'^END OF SECTION\s+([\d\.]+)\s*$', text_upper)
    if end_of_section_match:
        section_num = end_of_section_match.group(1)
        return (f"第 {section_num} 节结束", True)  # 已翻译

    # 去掉 "SECTION " 前缀
    # "SECTION 070150.19 - PREPARATION FOR REROOFING" → "070150.19 - PREPARATION FOR REROOFING"
    if text_upper.startswith("SECTION "):
        text = re.sub(r'^SECTION\s+', '', text, flags=re.IGNORECASE)

    # PART要翻译，不要删除！

    # 如果是列表样式（PR2, PR3等），去掉文本开头的编号
    # "2. Temporary protection..." → "Temporary protection..."
    if style_name and any(s in style_name.upper() for s in ['PR', 'LIST', 'CMT']):
        # 匹配开头的 "数字. " 或 "数字) "
        text = re.sub(r'^\d+[\.\)]\s+', '', text)

    # 删除英制单位表达式
    # 例如: "3 lb/100 sq. ft." "15 lb/100 sq. ft." 等
    imperial_patterns = [
        r'\d+(?:\.\d+)?\s*lb/\d+(?:\.\d+)?\s*sq\.\s*ft\.?',  # 如: 3 lb/100 sq. ft.
        r'\d+(?:\.\d+)?\s*lb/\d+(?:\.\d+)?\s*sq\s+ft\.?',    # 如: 3 lb/100 sq ft
        r'\d+(?:\.\d+)?\s*oz/\d+(?:\.\d+)?\s*sq\.\s*ft\.?',  # 盎司
        r'\d+(?:\.\d+)?\s*psi',                                # 磅力每平方英寸
        r'\d+(?:\.\d+)?\s*°F',                                 # 华氏度
    ]

    for pattern in imperial_patterns:
        # 删除英制单位，但保留前后的空格结构
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    # 清理多余的逗号和空格
    text = re.sub(r',\s*,', ',', text)  # 双逗号 → 单逗号
    text = re.sub(r'\s+,', ',', text)   # 空格+逗号 → 逗号
    text = re.sub(r',\s+\(', ' (', text) # 逗号+空格+括号 → 空格+括号
    text = re.sub(r'\s{2,}', ' ', text)  # 多个空格 → 单空格

    # 改进的占位符处理：简化为中文标记，让整句一起翻译
    # 例如 "A. Architect: <Insert name>." -> "A. Architect: 〈姓名〉." -> 翻译后 "A. 建筑师：〈姓名〉。"
    def replace_placeholder(match):
        content = match.group(1).strip()
        # 简化占位符为简短中文
        simple_placeholders = {
            'name': '姓名',
            'Name': '姓名',
            'license #': '证号',
            'license': '证书',
            'load': '荷载',
            'manufacturer': '制造商',
        }
        for key, value in simple_placeholders.items():
            if key.lower() in content.lower():
                return f'〈{value}〉'
        # 默认：简化为"待填"
        return '〈待填〉'

    # 替换占位符但保留在句子中
    text = re.sub(r'<Insert\s+([^>]+)>', replace_placeholder, text, flags=re.IGNORECASE)
    text = re.sub(r'<([^>]{1,50})>', replace_placeholder, text)

    return (text, False)  # 未翻译，需要送VLLM


def postprocess_translation(text: str, style_name: str = None) -> str:
    """后处理翻译结果（移除VLLM可能添加的编号）"""
    import re

    # 如果是列表样式（PR2, PR3, CMT等），去掉文本开头的编号
    # 这是为了防止VLLM忽略提示词，仍然添加编号
    if style_name and any(s in style_name.upper() for s in ['PR', 'LIST', 'CMT']):
        # 移除中文编号格式：1. 2. 3. 或 1、2、3、
        text = re.sub(r'^\d+[\.\)、]\s+', '', text)

    return text.strip()


def protect_units(text: str) -> tuple[str, dict]:
    """保护单位和列表编号不被翻译，返回（替换后的文本，映射表）"""
    import re

    units_map = {}
    counter = 0

    # ⚠️ 首先保护ASTM标准编号（必须在单位模式之前！）
    # 格式1: ASTM D4601/D4601M (带斜杠的完整格式)
    # 格式2: ASTM D4601M 或 ASTM C1177M (单独的公制版本)
    # 格式3: ASTM D4601 (标准版本)
    # 这样可以避免M后缀被误认为米(meter)单位
    astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'
    for match in re.finditer(astm_pattern, text):
        placeholder = f'⟨⟨ASTM_{counter}⟩⟩'
        units_map[placeholder] = match.group(0)
        text = text.replace(match.group(0), placeholder, 1)
        counter += 1

    # 定义需要保护的单位模式（⚠️ 只保护公制单位！）
    # ⚠️ 重要：复合单位必须在简单单位之前！
    unit_patterns = [
        # 公制复合单位（必须在前面，否则会被部分匹配）
        r'\d+(?:\.\d+)?\s*kg/sq\.\s*m',     # 面密度（英制写法，如 0.16 kg/sq. m）
        r'\d+(?:\.\d+)?\s*kg/cu\.\s*m',     # 体积密度（英制写法）
        r'\d+(?:\.\d+)?\s*kg/m²',           # 面密度（公制符号）
        r'\d+(?:\.\d+)?\s*kg/m³',           # 体积密度（公制符号）
        r'\d+(?:\.\d+)?\s*N/mm²',           # 压强
        r'\d+(?:\.\d+)?\s*W/m²',            # 功率密度
        r'\d+(?:\.\d+)?\s*mm²',             # 面积
        r'\d+(?:\.\d+)?\s*m²',              # 面积
        r'\d+(?:\.\d+)?\s*m³',              # 体积

        # 公制基本单位（在复合单位之后，按长度从长到短）
        r'\d+(?:\.\d+)?\s*(?:kPa|MPa)',     # 压强（3-4字母）
        r'\d+(?:\.\d+)?\s*(?:mm|cm|km|ml|Pa)(?![²³/a-z])',  # 2字母公制单位
        r'\d+(?:\.\d+)?\s*°C',              # 温度: °C
        r'\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])',  # 1字母公制单位（最后，防止部分匹配）
    ]

    # 保护列表编号（必须在行首或段落开头）
    # 格式: "a. " "1. " "A. " "I. " "(1) " 等
    list_number_patterns = [
        r'^([a-z]\.\s+)',         # a. b. c. 等
        r'^([A-Z]\.\s+)',         # A. B. C. 等
        r'^(\d+\.\s+)',           # 1. 2. 3. 等
        r'^([ivxlcdm]+\.\s+)',    # i. ii. iii. 等（罗马数字小写）
        r'^([IVXLCDM]+\.\s+)',    # I. II. III. 等（罗马数字大写）
        r'^(\(\d+\)\s+)',         # (1) (2) (3) 等
        r'^(\([a-z]\)\s+)',       # (a) (b) (c) 等
        r'^(\([A-Z]\)\s+)',       # (A) (B) (C) 等
    ]

    def replace_marker(match):
        nonlocal counter
        marker_text = match.group(0)
        # 使用更特殊的格式，避免被翻译API修改
        placeholder = f"⟨⟨UNIT{counter:04d}⟩⟩"
        units_map[placeholder] = marker_text
        counter += 1
        return placeholder

    result = text

    # 先保护列表编号（在行首，使用MULTILINE模式）
    for pattern in list_number_patterns:
        result = re.sub(pattern, replace_marker, result, flags=re.MULTILINE)

    # 再保护单位
    for pattern in unit_patterns:
        result = re.sub(pattern, replace_marker, result, flags=re.IGNORECASE)

    return result, units_map

def restore_units(text: str, units_map: dict) -> str:
    """恢复被保护的单位，并转换为符号格式"""
    import re

    result = text
    for placeholder, unit in units_map.items():
        # 转换英文单位为符号格式
        converted_unit = unit

        # kg/sq. m → kg/m² (处理非断空格 \xa0)
        converted_unit = re.sub(r'kg/sq\.\s*m', 'kg/m²', converted_unit, flags=re.IGNORECASE)
        # kg/cu. m → kg/m³
        converted_unit = re.sub(r'kg/cu\.\s*m', 'kg/m³', converted_unit, flags=re.IGNORECASE)
        # lb/sq. ft → lb/ft²
        converted_unit = re.sub(r'lb/sq\.\s*ft', 'lb/ft²', converted_unit, flags=re.IGNORECASE)
        # sq. ft → ft²
        converted_unit = re.sub(r'sq\.\s*ft', 'ft²', converted_unit, flags=re.IGNORECASE)
        # cu. m → m³
        converted_unit = re.sub(r'cu\.\s*m', 'm³', converted_unit, flags=re.IGNORECASE)

        result = result.replace(placeholder, converted_unit)
    return result

def apply_translation_with_colors(para, translated_text, unit_colors):
    """应用翻译到段落，并恢复单位的颜色（公制单位自动设置为蓝绿色）"""
    import re
    from docx.shared import RGBColor

    # 定义蓝绿色（Teal）
    TEAL_COLOR = RGBColor(0, 176, 176)  # 蓝绿色 #00B0B0

    # 检查翻译后的文本中是否包含公制单位
    # 注意：ASTM标准编号已在protect_units()中保护，不会被匹配
    unit_pattern = r'\d+(?:\.\d+)?\s*(?:kg/sq\.\s*m|kg/cu\.\s*m|kg/m²|kg/m³|N/mm²|W/m²|mm²|m²|m³|kPa|MPa|mm|cm|km|ml|Pa|°C)(?![²³/a-z])|\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])'
    has_units = bool(re.search(unit_pattern, translated_text, re.IGNORECASE))

    # 如果没有单位，按原来的方式处理
    if not has_units:
        if len(para.runs) == 1:
            original_run = para.runs[0]
            original_bold = original_run.bold
            original_italic = original_run.italic
            original_underline = original_run.underline
            original_font_size = original_run.font.size
            original_font_color = original_run.font.color.rgb if original_run.font.color and original_run.font.color.rgb else None

            original_run.text = translated_text

            if original_bold is not None:
                original_run.bold = original_bold
            if original_italic is not None:
                original_run.italic = original_italic
            if original_underline is not None:
                original_run.underline = original_underline
            if original_font_size:
                original_run.font.size = original_font_size
            if original_font_color:
                original_run.font.color.rgb = original_font_color
            original_run.font.name = '宋体'
        else:
            # 多个run，简化处理
            first_run = para.runs[0]
            first_run.text = translated_text
            for run in para.runs[1:]:
                run.text = ''
        return

    # 有单位颜色，需要拆分并应用颜色
    # ⚠️ 重要：在拆分文本之前，先临时保护ASTM标准编号！
    # 因为此时translated_text中的ASTM占位符已经被restore_units()恢复了
    # 如果不保护，会匹配到ASTM编号中的M后缀
    astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'
    astm_placeholders = {}
    astm_counter = 0
    temp_text = translated_text
    for match in re.finditer(astm_pattern, temp_text):
        placeholder = f'⟨⟨TEMP_ASTM_{astm_counter}⟩⟩'
        astm_placeholders[placeholder] = match.group(0)
        temp_text = temp_text.replace(match.group(0), placeholder, 1)
        astm_counter += 1

    # 找到所有公制单位的位置（必须与protect_units()中的顺序保持一致）
    unit_pattern = r'\d+(?:\.\d+)?\s*(?:kg/sq\.\s*m|kg/cu\.\s*m|kg/m²|kg/m³|N/mm²|W/m²|mm²|m²|m³|kPa|MPa|mm|cm|km|ml|Pa|°C)(?![²³/a-z])|\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])'

    # 分割文本，保留单位（使用临时保护后的文本）
    parts = re.split(f'({unit_pattern})', temp_text, flags=re.IGNORECASE)

    # 保存第一个run的格式
    if para.runs:
        first_run = para.runs[0]
        default_bold = first_run.bold
        default_italic = first_run.italic
        default_underline = first_run.underline
        default_font_size = first_run.font.size
        default_font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
    else:
        default_bold = None
        default_italic = None
        default_underline = None
        default_font_size = None
        default_font_color = None

    # 清空所有runs
    for run in para.runs[:]:
        run._element.getparent().remove(run._element)

    # 重建runs
    for part in parts:
        if not part:
            continue

        # ⚠️ 恢复ASTM占位符
        final_part = part
        for placeholder, astm_text in astm_placeholders.items():
            final_part = final_part.replace(placeholder, astm_text)

        run = para.add_run(final_part)

        # 检查是否是单位（在占位符替换前检查）
        is_unit = bool(re.match(unit_pattern, part, re.IGNORECASE))

        if is_unit:
            # 公制单位始终使用蓝绿色
            run.font.color.rgb = TEAL_COLOR
        else:
            # 非单位部分，应用默认格式
            if default_font_color:
                run.font.color.rgb = default_font_color

        # 应用其他格式
        if default_bold is not None:
            run.bold = default_bold
        if default_italic is not None:
            run.italic = default_italic
        if default_underline is not None:
            run.underline = default_underline
        if default_font_size:
            run.font.size = default_font_size
        run.font.name = '宋体'

def should_skip_by_style(style_name: str) -> bool:
    """根据样式判断是否删除（空的CMT/TIP段落）"""
    skip_styles = ["CMT", "TIP"]
    return style_name in skip_styles

# ============ LibreOffice 转换 ============
def convert_doc_to_docx(input_file: Path, output_file: Path) -> bool:
    """用 LibreOffice 转换 DOC → DOCX (使用临时用户配置避免Java问题)"""
    import tempfile
    import shutil

    try:
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # 创建临时用户配置文件目录
        user_profile = tempfile.mkdtemp(prefix='libreoffice_')

        try:
            result = subprocess.run([
                LIBREOFFICE,
                '--headless',
                '--invisible',
                '--nocrashreport',
                '--nodefault',
                '--nofirststartwizard',
                '--nolockcheck',
                '--nologo',
                '--norestore',
                f'-env:UserInstallation=file://{user_profile}',
                '--convert-to', 'docx:MS Word 2007 XML',
                '--outdir', str(output_file.parent),
                str(input_file)
            ], capture_output=True, text=True, timeout=60)

            # LibreOffice 会生成同名的 .docx 文件
            generated_file = output_file.parent / (input_file.stem + '.docx')

            if generated_file.exists() and generated_file != output_file:
                # 重命名为目标文件名
                generated_file.rename(output_file)
                return True
            elif output_file.exists():
                return True
            else:
                print(f"  ❌ 转换失败: {result.stderr}")
                return False
        finally:
            # 清理临时目录
            try:
                shutil.rmtree(user_profile)
            except:
                pass

    except Exception as e:
        print(f"  ❌ 转换错误: {e}")
        return False

# ============ 翻译API ============
def translate_single(text: str) -> str:
    if not text.strip():
        return text

    prompt = f"""将以下英文完整翻译成中文，不要遗漏任何单词：

{text}

翻译结果："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,  # 降低温度以获得更确定的翻译
                "max_tokens": 2000
            },
            timeout=60
        )

        if response.status_code == 200:
            result = response.json()['choices'][0]['message']['content'].strip()

            # 移除thinking标签
            if '<think>' in result:
                if '</think>' in result:
                    result = result.split('</think>')[-1].strip()
                else:
                    return text

            # 移除前缀
            prefixes = ["中文：", "翻译：", "翻译结果：", "译文："]
            for prefix in prefixes:
                if result.startswith(prefix):
                    result = result.split("：", 1)[1].strip()
                    break

            # 只取第一行（避免模型输出多行解释）
            result = result.split('\n')[0].strip()

            return result if result else text
        return text
    except Exception as e:
        print(f"    ⚠️ 翻译失败: {e}")
        return text

def remove_imperial_units(text: str) -> str:
    """删除英制单位（保留括号中的公制单位）

    处理各种格式：
    - 整数: "9 inch (230 mm)"
    - 分数: "1/4 inch (6 mm)", "15/32 inch (12 mm)"
    - 带连字符: "1-inch- (25-mm-)", "15/32-inch (12-mm)"
    """
    import re

    # 数字模式：支持整数、小数、分数（如 1/4, 15/32）
    # 示例: "1", "1.5", "1/4", "15/32"
    num_pattern = r'\d+(?:/\d+)?(?:\.\d+)?'

    # 策略：匹配"英制单位 (公制单位)"这种模式，只保留括号部分
    # 比如 "9 inch (230 mm)" → "(230 mm)"
    # 比如 "1/4 inch (6 mm)" → "(6 mm)"
    # 比如 "1-inch- (25-mm-)" → "(25-mm-)"

    # 模式1: 数字 + 英制单位 + (公制单位) → 只保留 (公制单位)
    patterns_with_metric = [
        # "9 inch (230 mm)" 类型，支持连字符和分数
        rf'{num_pattern}\s*-?\s*(?:inch|inches|in\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:foot|feet|ft\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:yard|yards|yd\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:mile|miles|mi\.?)\s*-?\s*(\([^)]+\))',

        # "3 lb/100 sq. ft. (1.5 kg/m²)" 类型
        rf'{num_pattern}\s*lbs?\.?/{num_pattern}\s*sq\.?\s*ft\.?\s*(\([^)]+\))',

        # "XX lb/cu. ft. (128 kg/cu. m)" 类型
        rf'{num_pattern}\s*lbs?\.?/cu\.?\s*ft\.?\s*(\([^)]+\))',

        # 其他复合单位
        rf'{num_pattern}\s*-?\s*(?:pound|pounds|lb\.?|lbs\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:sq\.?\s*ft\.?|square\s+feet?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:sq\.?\s*in\.?|square\s+inches?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:cu\.?\s*ft\.?|cubic\s+feet?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:ounce|ounces|oz\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:gallon|gallons|gal\.?)\s*-?\s*(\([^)]+\))',
    ]

    result = text

    # 替换为只保留括号部分
    for pattern in patterns_with_metric:
        result = re.sub(pattern, r'\1', result, flags=re.IGNORECASE)

    # 模式2: 删除没有公制单位的英制单位
    # （如果后面没有括号，就直接删除）
    # 注意：长单词在前（inches before inch），避免部分匹配
    standalone_imperial = [
        rf'{num_pattern}\s*-?\s*(?:inches|inch|in\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:feet|foot|ft\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:yards|yard|yd\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:miles|mile|mi\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:pounds|pound|lbs?\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:ounces|ounce|oz\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:square\s+feet|square\s+foot|sq\.?\s*ft\.?)(?!\s*\()',
        rf'{num_pattern}\s*lbs?\.?/{num_pattern}\s*sq\.?\s*ft\.?(?!\s*\()',
    ]

    for pattern in standalone_imperial:
        result = re.sub(pattern, '', result, flags=re.IGNORECASE)

    # 模式3: 清理残留的单位（比如前面数字被删除后剩下的 "/cu. ft."）
    leftover_units = [
        r'/\s*cu\.?\s*ft\.?',  # /cu. ft.
        r'/\s*sq\.?\s*ft\.?',  # /sq. ft.
        r'/\s*cu\.?\s*in\.?',  # /cu. in.
        r'/\s*sq\.?\s*in\.?',  # /sq. in.
    ]

    for pattern in leftover_units:
        result = re.sub(pattern, '', result, flags=re.IGNORECASE)

    # 清理多余的空格和标点
    result = re.sub(r'\s+', ' ', result)  # 多个空格变一个
    result = re.sub(r'\s*,\s*,', ',', result)  # 删除多余的逗号
    result = re.sub(r'\s*:\s*\(', ' (', result)  # "Thickness: (xxx)" → "Thickness (xxx)"
    result = re.sub(r'\[\s*/\s*\(', '[(', result)  # "[/ (" → "[("
    result = result.strip()

    return result

def translate_placeholders(text: str) -> str:
    """只翻译占位符（不删除英制单位）"""
    replacements = {
        "<Insert name>": "<插入姓名>",
        "<Insert Name>": "<插入姓名>",
        "<insert name>": "<插入姓名>",
        "<Insert license #>": "<插入许可证编号>",
        "<Insert License #>": "<插入许可证编号>",
        "<Insert list of Sections>": "<插入章节列表>",
        "<Insert List of Sections>": "<插入章节列表>",
        "<Insert section list>": "<插入章节列表>",
        "<Insert address>": "<插入地址>",
        "<Insert Address>": "<插入地址>",
        "<Insert phone>": "<插入电话>",
        "<Insert Phone>": "<插入电话>",
        "<Insert email>": "<插入邮箱>",
        "<Insert Email>": "<插入邮箱>",
        "<Insert date>": "<插入日期>",
        "<Insert Date>": "<插入日期>",
        "<Insert project name>": "<插入项目名称>",
        "<Insert Project Name>": "<插入项目名称>",
        "<Insert description>": "<插入描述>",
        "<Insert Description>": "<插入描述>",
        "<Insert value>": "<插入数值>",
        "<Insert Value>": "<插入数值>",
        "<Insert location>": "<插入位置>",
        "<Insert Location>": "<插入位置>",
        "<Insert area>": "<插入面积>",
        "<Insert Area>": "<插入面积>",
        "<Insert dimension>": "<插入尺寸>",
        "<Insert Dimension>": "<插入尺寸>",
    }

    result = text
    for eng, chn in replacements.items():
        result = result.replace(eng, chn)

    return result

def translate_batch(texts: List[str]) -> List[str]:
    if not texts:
        return []

    non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
    non_empty_texts = [texts[i] for i in non_empty_indices]

    if not non_empty_texts:
        return texts

    combined = "\n---SPLIT---\n".join(non_empty_texts)
    prompt = f"""你是建筑工程规范翻译专家。将以下英文段落翻译成中文（段落之间用---SPLIT---分隔）：

{combined}

翻译要求：
1. **必须翻译所有文本**，包括全大写的章节标题（如"ACTION SUBMITTALS"翻译为"操作提交件"）
2. 保留专业缩写（如ASTM、EPS、OSB保持不变）
3. **不要添加任何数字编号**，保持原文格式
4. 用---SPLIT---分隔每个翻译结果
5. 只输出翻译，不要添加解释

中文翻译："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": len(combined) * 3
            },
            timeout=600
        )

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            if '<think>' in translated:
                if '</think>' in translated:
                    translated = translated.split('</think>')[-1].strip()
                else:
                    return [translate_single(t) for t in texts]

            if translated.startswith("翻译：") or translated.startswith("中文翻译："):
                translated = translated.split("：", 1)[1].strip()

            parts = translated.split("---SPLIT---")

            if len(parts) == len(non_empty_texts):
                results = texts.copy()
                for i, idx in enumerate(non_empty_indices):
                    results[idx] = parts[i].strip()
                return results
            else:
                return [translate_single(t) for t in texts]
        else:
            return texts
    except Exception as e:
        return [translate_single(t) for t in texts]

# ============ 翻译DOCX文档 ============
def translate_docx(docx_file: Path) -> bool:
    """翻译DOCX文档（完全保留样式）"""
    try:
        from docx import Document

        doc = Document(str(docx_file))

        # 移除页眉页脚（完全删除所有XML内容包括表格）
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # 彻底清空页眉（包括段落和表格）
            section.header._element.clear()
            section.first_page_header._element.clear()
            section.even_page_header._element.clear()

            # 彻底清空页脚（包括段落和表格）
            section.footer._element.clear()
            section.first_page_footer._element.clear()
            section.even_page_footer._element.clear()

        # ========== 删除所有文本框（包括TIP框）==========
        print("  🗑️  删除文本框...")
        deleted_textbox_count = 0
        try:
            # 方法1: 删除段落中的文本框
            paragraphs_with_textbox = []
            for para in doc.paragraphs:
                para_xml = para._element.xml
                if b'<w:txbxContent>' in para_xml or b'<v:textbox>' in para_xml:
                    paragraphs_with_textbox.append(para)

            for para in paragraphs_with_textbox:
                p_element = para._element
                p_element.getparent().remove(p_element)
                deleted_textbox_count += 1

            # 方法2: 删除body中的独立shape对象
            body = doc._element.body

            # 删除现代格式的文本框 (drawing with txbx)
            for drawing in body.findall('.//w:drawing', namespaces=body.nsmap):
                if drawing.find('.//wps:txbx', namespaces=body.nsmap) is not None:
                    parent = drawing.getparent()
                    if parent is not None:
                        parent.remove(drawing)
                        deleted_textbox_count += 1

            # 删除VML格式的文本框（旧版Word）
            for shape in body.findall('.//v:shape', namespaces={'v': 'urn:schemas-microsoft-com:vml'}):
                if shape.find('.//v:textbox', namespaces={'v': 'urn:schemas-microsoft-com:vml'}) is not None:
                    parent = shape.getparent()
                    if parent is not None:
                        parent.remove(shape)
                        deleted_textbox_count += 1

            print(f"  ✅ 已删除 {deleted_textbox_count} 个文本框")
        except Exception as e:
            print(f"  ⚠️  文本框删除警告: {e}")

        # 第一步：标记并删除需要跳过的段落
        paragraphs_to_delete = []
        skipped_count = 0

        for para in doc.paragraphs:
            # 删除条件1：包含跳过关键词的段落
            if para.text.strip() and should_skip_paragraph(para.text):
                paragraphs_to_delete.append(para)
                skipped_count += 1
            # 删除条件2：空的CMT/TIP段落
            elif not para.text.strip() and should_skip_by_style(para.style.name):
                paragraphs_to_delete.append(para)
                skipped_count += 1

        # 删除标记的段落
        for para in paragraphs_to_delete:
            p_element = para._element
            p_element.getparent().remove(p_element)

        # 第二步：收集需要翻译的段落（同时记录单位的颜色）
        paragraphs_to_translate = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 扫描run，找到包含单位的颜色
                unit_colors = {}  # {单位文本: RGB颜色}
                for run in para.runs:
                    if run.text and run.font.color and run.font.color.rgb:
                        # 检查是否包含单位模式（与protect_units保持一致）
                        import re
                        unit_pattern = r'\d+(?:\.\d+)?\s*(?:kg/sq\.\s*m|kg/cu\.\s*m|lb/sq\.\s*ft|kg/m²|kg/m³|N/mm²|W/m²|mm²|m²|m³|sq\.\s*ft|cu\.\s*m|mm|cm|m|km|ft|in|kg|g|t|lb|ml|l|L|°C|Pa|kPa|MPa|kW)'
                        units = re.findall(unit_pattern, run.text, re.IGNORECASE)
                        for unit in units:
                            unit_colors[unit.strip()] = run.font.color.rgb

                # 保存段落格式（用于后续恢复）
                para_format = {
                    'space_before': para.paragraph_format.space_before,
                    'space_after': para.paragraph_format.space_after,
                    'line_spacing': para.paragraph_format.line_spacing,
                    'line_spacing_rule': para.paragraph_format.line_spacing_rule
                }

                paragraphs_to_translate.append({
                    'paragraph': para,
                    'unit_colors': unit_colors,
                    'format': para_format
                })

        total = len(paragraphs_to_translate)
        print(f"  📊 共需翻译 {total} 个段落")

        # 批量翻译
        for i in range(0, total, BATCH_SIZE):
            current_batch = min(i + BATCH_SIZE, total)
            progress = f"[{current_batch}/{total}]"
            print(f"  🔄 {progress} 正在翻译...")

            batch_data = paragraphs_to_translate[i:i+BATCH_SIZE]
            batch_paras = [item['paragraph'] for item in batch_data]
            batch_texts = [p.text for p in batch_paras]

            # 第1步：预处理（去掉SECTION前缀，去掉列表编号等）
            # 返回 (文本, 是否已翻译)
            preprocess_results = [preprocess_text(text, para.style.name if para.style else None)
                                 for text, para in zip(batch_texts, batch_paras)]
            batch_texts = [result[0] for result in preprocess_results]
            already_translated = [result[1] for result in preprocess_results]

            # 第2步：删除英制单位（在翻译之前）
            batch_texts = [remove_imperial_units(t) for t in batch_texts]

            # 第3步：保护单位不被翻译
            protected_texts = []
            units_maps = []
            for text in batch_texts:
                protected_text, units_map = protect_units(text)
                protected_texts.append(protected_text)
                units_maps.append(units_map)

            # 第4步：翻译（跳过已翻译的）
            translated_texts = []
            texts_to_translate = []
            indices_to_translate = []

            for idx, (text, is_translated) in enumerate(zip(protected_texts, already_translated)):
                if is_translated:
                    # 已经是中文，不需要翻译
                    translated_texts.append(text)
                else:
                    # 需要翻译
                    texts_to_translate.append(text)
                    indices_to_translate.append(idx)
                    translated_texts.append(None)  # 占位

            # 只翻译需要翻译的文本
            if texts_to_translate:
                vllm_translations = translate_batch(texts_to_translate)
                # 填回翻译结果
                for idx, trans in zip(indices_to_translate, vllm_translations):
                    translated_texts[idx] = trans

            # 第5步：恢复单位
            translated_texts = [restore_units(trans, units_map)
                              for trans, units_map in zip(translated_texts, units_maps)]

            # 应用翻译（保留所有格式包括颜色）
            for data_item, para, trans_text in zip(batch_data, batch_paras, translated_texts):
                # 第6步：后处理：翻译占位符
                trans_text = translate_placeholders(trans_text)

                # 第7步：后处理：移除VLLM可能添加的编号
                trans_text = postprocess_translation(trans_text, para.style.name if para.style else None)

                # 使用新函数应用翻译并恢复单位颜色
                apply_translation_with_colors(para, trans_text, data_item['unit_colors'])

                # 恢复段落格式（如果LibreOffice转换时丢失了格式，使用默认值）
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING

                saved_format = data_item['format']
                # 如果原始格式为None，使用默认值（12pt段前，1.0行距）
                if saved_format['space_before'] is not None:
                    para.paragraph_format.space_before = saved_format['space_before']
                else:
                    para.paragraph_format.space_before = Pt(12)

                if saved_format['space_after'] is not None:
                    para.paragraph_format.space_after = saved_format['space_after']

                if saved_format['line_spacing'] is not None and saved_format['line_spacing_rule'] is not None:
                    para.paragraph_format.line_spacing = saved_format['line_spacing']
                    para.paragraph_format.line_spacing_rule = saved_format['line_spacing_rule']
                else:
                    # 设置单倍行距（1.0）
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 翻译表格
        if doc.tables:
            for table in doc.tables:
                cells_data = []
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                cells_data.append((para, para.text))

                if cells_data:
                    cell_texts = [text for _, text in cells_data]

                    # 保护单位
                    protected_texts = []
                    units_maps = []
                    for text in cell_texts:
                        protected_text, units_map = protect_units(text)
                        protected_texts.append(protected_text)
                        units_maps.append(units_map)

                    # 翻译
                    translated = translate_batch(protected_texts)

                    # 恢复单位
                    translated = [restore_units(trans, units_map)
                                for trans, units_map in zip(translated, units_maps)]

                    for (para, _), trans in zip(cells_data, translated):
                        # 后处理：翻译占位符
                        trans = translate_placeholders(trans)

                        if len(para.runs) == 1:
                            # 保存原有格式
                            original_run = para.runs[0]
                            original_bold = original_run.bold
                            original_italic = original_run.italic
                            original_underline = original_run.underline
                            original_font_size = original_run.font.size
                            original_font_color = original_run.font.color.rgb if original_run.font.color and original_run.font.color.rgb else None

                            # 修改文本
                            original_run.text = trans

                            # 恢复格式（包括颜色）
                            if original_bold is not None:
                                original_run.bold = original_bold
                            if original_italic is not None:
                                original_run.italic = original_italic
                            if original_underline is not None:
                                original_run.underline = original_underline
                            if original_font_size:
                                original_run.font.size = original_font_size
                            if original_font_color:
                                original_run.font.color.rgb = original_font_color
                            original_run.font.name = '宋体'

                        elif len(para.runs) > 1:
                            # 保存第一个run的格式
                            first_run = para.runs[0]
                            original_bold = first_run.bold
                            original_italic = first_run.italic
                            original_underline = first_run.underline
                            original_font_size = first_run.font.size
                            original_font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None

                            # 修改文本
                            first_run.text = trans

                            # 恢复格式
                            if original_bold is not None:
                                first_run.bold = original_bold
                            if original_italic is not None:
                                first_run.italic = original_italic
                            if original_underline is not None:
                                first_run.underline = original_underline
                            if original_font_size:
                                first_run.font.size = original_font_size
                            if original_font_color:
                                first_run.font.color.rgb = original_font_color
                            first_run.font.name = '宋体'

                            # 删除其他runs
                            for run in para.runs[1:]:
                                run.text = ''
                        else:
                            new_run = para.add_run(trans)
                            new_run.font.name = '宋体'

        # 保存
        doc.save(str(docx_file))
        return True

    except Exception as e:
        print(f"  ❌ 翻译错误: {e}")
        return False

# ============ 批量处理 ============
def batch_process(max_files=None):
    source_path = Path(SOURCE_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    folders = sorted([d for d in source_path.iterdir() if d.is_dir() and not d.name.startswith('.')])

    print(f"\n{'='*70}")
    print(f"  🚀 批量转换+翻译（LibreOffice + 完整样式保留）")
    print(f"{'='*70}")
    print(f"步骤1: LibreOffice 转换 DOC → DOCX")
    print(f"步骤2: 翻译 DOCX（保留所有样式）")
    print(f"文件夹: {len(folders)} 个")
    print(f"{'='*70}\n")

    total_success = 0
    total_fail = 0
    total_skip = 0

    for folder_idx, folder in enumerate(folders, 1):
        folder_name = folder.name
        translated_folder_name = translate_folder_name(folder_name)

        print(f"\n{'#'*70}")
        print(f"# [{folder_idx}/{len(folders)}] {folder_name}")
        print(f"# → {translated_folder_name}")
        print(f"{'#'*70}")

        files = sorted(list(folder.glob("*.DOC")) + list(folder.glob("*.docx")))
        files = [f for f in files if not f.name.startswith('~$') and not f.name.startswith('.~')]

        if max_files:
            files = files[:max_files]

        for file_idx, input_file in enumerate(files, 1):
            file_name = input_file.name
            translated_file_name = translate_file_name(file_name)

            output_folder = Path(OUTPUT_DIR).parent / translated_folder_name
            output_file = output_folder / translated_file_name

            print(f"\n  [{file_idx}/{len(files)}] {file_name}")
            print(f"       → {translated_file_name}")

            if output_file.exists():
                print(f"  ⏭️  已存在，跳过")
                total_skip += 1
                continue

            try:
                # 步骤1：转换
                print(f"  🔄 转换中...", end='', flush=True)
                if convert_doc_to_docx(input_file, output_file):
                    print(f" ✅")

                    # 步骤2：翻译
                    print(f"  📝 翻译中...", end='', flush=True)
                    if translate_docx(output_file):
                        print(f" ✅")
                        total_success += 1
                    else:
                        print(f" ❌")
                        total_fail += 1
                else:
                    print(f" ❌")
                    total_fail += 1

            except KeyboardInterrupt:
                print("\n\n⚠️  用户中断")
                break
            except Exception as e:
                print(f"\n  ❌ 失败: {e}")
                total_fail += 1

        if max_files:
            break

    print(f"\n\n{'='*70}")
    print(f"  📊 处理完成")
    print(f"{'='*70}")
    print(f"✅ 成功: {total_success}")
    print(f"⏭️  跳过: {total_skip}")
    print(f"❌ 失败: {total_fail}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="批量转换+翻译")
    parser.add_argument("--test", action="store_true", help="测试（1个文件）")
    parser.add_argument("--small", action="store_true", help="小批量（10个文件）")

    args = parser.parse_args()

    print("🔌 测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM连接成功\n")
        else:
            print(f"⚠️  VLLM响应异常")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        sys.exit(1)

    max_files = None
    if args.test:
        max_files = 1
    elif args.small:
        max_files = 10

    batch_process(max_files)
