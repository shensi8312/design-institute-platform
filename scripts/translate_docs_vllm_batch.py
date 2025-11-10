#!/usr/bin/env python3
"""
批量文档翻译脚本 - 批量优化版
每次翻译10个段落，速度提升10倍+
"""

import os
import sys
import subprocess
from pathlib import Path
from typing import List
import requests

sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

SOURCE_DIR = "docs/specs"
OUTPUT_DIR = "docs/specs_zh"
VLLM_URL = os.getenv("VLLM_URL", "http://10.10.18.3:8000")
VLLM_MODEL = os.getenv("VLLM_MODEL", "/mnt/data/models/Qwen3-32B")
BATCH_SIZE = 10

def translate_batch(texts: List[str]) -> List[str]:
    """批量翻译"""
    if not texts:
        return []

    combined = "\n---SPLIT---\n".join(texts)

    prompt = f"""你是建筑工程规范翻译专家。翻译以下{len(texts)}个段落（用---SPLIT---分隔）：

{combined}

要求：保留专业术语、格式、标记。用---SPLIT---分隔翻译结果。

中文翻译："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": len(combined) * 3
            },
            timeout=180
        )

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            if translated.startswith("翻译："):
                translated = translated[3:].strip()
            if translated.startswith("中文翻译："):
                translated = translated[5:].strip()

            parts = translated.split("---SPLIT---")

            if len(parts) == len(texts):
                return [p.strip() for p in parts]
            else:
                print(f"  ⚠️  分割失败({len(parts)}!={len(texts)})，逐个翻译")
                return [translate_single(t) for t in texts]
        else:
            return texts
    except Exception as e:
        print(f"  ⚠️  批量失败: {e}")
        return [translate_single(t) for t in texts]

def translate_single(text: str) -> str:
    """单个翻译"""
    if not text.strip():
        return text

    prompt = f"""翻译建筑规范文本（保留术语和格式）：

{text}

中文："""

    try:
        response = requests.post(
            f"{VLLM_URL}/v1/chat/completions",
            json={
                "model": VLLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 1000
            },
            timeout=60
        )

        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content'].strip()
        return text
    except:
        return text

def process_doc(input_file: Path, output_file: Path):
    """处理文档"""
    print(f"\n{'='*60}")
    print(f"处理: {input_file.name}")
    print(f"{'='*60}")

    # DOC转DOCX
    if input_file.suffix.upper() == '.DOC':
        temp_docx = Path(f"/tmp/{input_file.stem}_temp.docx")
        if sys.platform == 'darwin':
            result = subprocess.run([
                'textutil', '-convert', 'docx',
                str(input_file), '-output', str(temp_docx)
            ], capture_output=True)

            if result.returncode != 0 or not temp_docx.exists():
                print(f"跳过: 无法转换")
                return
            working_file = temp_docx
        else:
            print("跳过: 非macOS系统")
            return
    else:
        working_file = input_file

    output_file.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document

        doc = Document(str(working_file))

        # 移除页眉页脚
        print("移除页眉页脚...")
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for p in section.header.paragraphs:
                p.clear()
            for p in section.footer.paragraphs:
                p.clear()

        # 收集段落
        paras = []
        formats = []
        for p in doc.paragraphs:
            if p.text.strip():
                paras.append(p)
                fmt = []
                for run in p.runs:
                    fmt.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })
                formats.append(fmt)

        total = len(paras)
        print(f"批量翻译 {total} 个段落（{BATCH_SIZE}个/批）...\n")

        # 批量翻译
        for i in range(0, total, BATCH_SIZE):
            batch = paras[i:i+BATCH_SIZE]
            batch_texts = [p.text for p in batch]

            print(f"[{i+1}-{min(i+BATCH_SIZE, total)}/{total}] 批量翻译中...")

            translated = translate_batch(batch_texts)

            # 应用翻译
            for j, (para, trans_text) in enumerate(zip(batch, translated)):
                para.clear()
                idx = i + j
                if formats[idx]:
                    run = para.add_run(trans_text)
                    fmt = formats[idx][0]
                    if fmt['bold'] is not None:
                        run.bold = fmt['bold']
                    if fmt['italic'] is not None:
                        run.italic = fmt['italic']
                    if fmt['underline'] is not None:
                        run.underline = fmt['underline']
                    if fmt['font_name']:
                        run.font.name = fmt['font_name']
                    if fmt['font_size']:
                        run.font.size = fmt['font_size']
                    if fmt['font_color']:
                        run.font.color.rgb = fmt['font_color']
                else:
                    para.add_run(trans_text)

        # 翻译表格
        if doc.tables:
            print(f"\n翻译 {len(doc.tables)} 个表格...")
            for table in doc.tables:
                cells = []
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                cells.append(para)

                if cells:
                    cell_texts = [c.text for c in cells]
                    translated = translate_batch(cell_texts)

                    for para, trans in zip(cells, translated):
                        bold = para.runs[0].bold if para.runs else None
                        italic = para.runs[0].italic if para.runs else None
                        para.clear()
                        run = para.add_run(trans)
                        run.bold = bold
                        run.italic = italic

        # 保存
        doc.save(str(output_file))
        print(f"\n✅ 完成！保存到: {output_file}")

        # 清理
        if input_file.suffix.upper() == '.DOC' and temp_docx.exists():
            temp_docx.unlink()

    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()

def batch_translate(max_files=None):
    """批量翻译所有文档"""
    source_path = Path(SOURCE_DIR)
    output_path = Path(OUTPUT_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    files = list(source_path.rglob("*.DOC")) + list(source_path.rglob("*.docx"))
    files = [f for f in files if not f.name.startswith('~$')]

    if max_files:
        files = files[:max_files]

    print(f"\n{'='*60}")
    print(f"  批量文档翻译（优化版）")
    print(f"{'='*60}")
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"VLLM: {VLLM_URL}")
    print(f"模型: {VLLM_MODEL}")
    print(f"批量: {BATCH_SIZE} 段/次")
    print(f"文件: {len(files)} 个")
    print(f"{'='*60}\n")

    success = 0
    fail = 0

    for i, input_file in enumerate(files, 1):
        print(f"\n\n{'#'*60}")
        print(f"# 进度: [{i}/{len(files)}]")
        print(f"{'#'*60}")

        rel_path = input_file.relative_to(source_path)
        output_file = output_path / rel_path

        if output_file.suffix.upper() == '.DOC':
            output_file = output_file.with_suffix('.docx')

        if output_file.exists():
            print(f"⏭️  已存在，跳过")
            continue

        try:
            process_doc(input_file, output_file)
            success += 1
        except KeyboardInterrupt:
            print("\n⚠️  用户中断")
            break
        except Exception as e:
            print(f"\n❌ 失败: {e}")
            fail += 1

    print(f"\n\n{'='*60}")
    print(f"  翻译完成")
    print(f"{'='*60}")
    print(f"✅ 成功: {success}")
    print(f"❌ 失败: {fail}")
    print(f"输出: {OUTPUT_DIR}")
    print(f"{'='*60}\n")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="批量文档翻译（优化版）")
    parser.add_argument("--test", action="store_true", help="测试（1个文件）")
    parser.add_argument("--small", action="store_true", help="小批量（10个文件）")

    args = parser.parse_args()

    # 测试连接
    print("测试VLLM连接...")
    try:
        response = requests.get(f"{VLLM_URL}/v1/models", timeout=5)
        if response.status_code == 200:
            print(f"✅ VLLM服务连接成功: {VLLM_URL}\n")
        else:
            print(f"⚠️  VLLM响应异常: {response.status_code}")
    except Exception as e:
        print(f"❌ VLLM连接失败: {e}")
        sys.exit(1)

    max_files = None
    if args.test:
        max_files = 1
    elif args.small:
        max_files = 10

    batch_translate(max_files)
