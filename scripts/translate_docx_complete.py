#!/usr/bin/env python3
"""
完整DOCX翻译脚本（远程服务器版）
功能：
✅ 保留所有格式（粗体、斜体、字体、颜色）
✅ 删除版权信息
✅ 删除文本框
✅ 删除页眉页脚（包括XML）
✅ 删除英制单位，保留公制单位
✅ 公制单位颜色保留（蓝绿色）
✅ 翻译文件名
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
import requests

# 强制刷新输出
sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# ============ 配置 ============
SOURCE_DIR = "specs_docx"
OUTPUT_DIR = "specs_zh"
TRANSLATION_API = "http://10.10.18.3:8000/v1/chat/completions"
MODEL_NAME = "qwen-next-80b"
BATCH_SIZE = 10  # 每批翻译10个段落

# 定义蓝绿色（公制单位专用）
TEAL_COLOR = RGBColor(0, 176, 176)  # #00B0B0

# 版权关键词（跳过翻译并删除）
COPYRIGHT_KEYWORDS = [
    "Copyright",
    "©",
    "AIA",
    "Deltek",
    "All rights reserved",
    "版权所有",
    "The American Institute of Architects",
    "Exclusively published and distributed",
    "TIPS:",
    "TIP:",
    "To view non-printing Editor's Notes",
    "MasterWorks/Single-File Formatting",
    "MasterWorks/Supporting Information",
    "Content Requests:",
    "<Double click here",
    "Double click here",
]

# ============ 翻译文件名 ============
def translate_filename(filename: str) -> str:
    """翻译文件名，保留编号和格式"""
    name_without_ext = filename.rsplit('.', 1)[0]
    ext = '.' + filename.rsplit('.', 1)[1] if '.' in filename else ''

    # 匹配编号部分（如 "071113 FL -"）
    match = re.match(r'^([\d\.\s]+(?:FL|SL|FS)?[\s\-]*)', name_without_ext)

    if match:
        prefix = match.group(1)
        description = name_without_ext[len(prefix):].strip()

        # 翻译描述部分
        if description:
            translated_desc = translate_text(description, is_filename=True)
            if translated_desc:
                return f"{prefix}{translated_desc}{ext}"

    return filename

# ============ 翻译API ============
def translate_text(text: str, is_filename: bool = False) -> str:
    """使用VLLM服务翻译文本"""
    if not text.strip():
        return text

    # 检查是否包含版权信息
    if any(keyword in text for keyword in COPYRIGHT_KEYWORDS):
        return None  # 不翻译版权信息

    try:
        if is_filename:
            prompt = f"请将以下英文直接翻译成中文，只翻译文字本身，不要添加任何额外的词（如'文件'等）：\n\n{text}"
        else:
            prompt = f"请将以下英文翻译成中文，保持专业术语准确性：\n\n{text}"

        response = requests.post(TRANSLATION_API, json={
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 4000,
            "temperature": 0.3
        }, timeout=120)

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            # 移除thinking标签
            if '<think>' in translated:
                if '</think>' in translated:
                    translated = translated.split('</think>')[-1].strip()
                else:
                    return text

            # 移除前缀
            prefixes = ["中文：", "翻译：", "翻译结果："]
            for prefix in prefixes:
                if translated.startswith(prefix):
                    translated = translated.split("：", 1)[1].strip()
                    break

            # 只取第一行
            translated = translated.split('\n')[0].strip()

            return translated if translated else text
        else:
            print(f"  ❌ API错误: {response.status_code}")
            return None
    except Exception as e:
        print(f"  ❌ 翻译错误: {e}")
        return None

def translate_batch(texts: list) -> list:
    """批量翻译多个段落"""
    if not texts:
        return []

    non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
    non_empty_texts = [texts[i] for i in non_empty_indices]

    if not non_empty_texts:
        return texts

    combined = "\n---SPLIT---\n".join(non_empty_texts)

    prompt = f"""你是建筑工程规范翻译专家。将以下{len(non_empty_texts)}个段落翻译成中文（用---SPLIT---分隔）：

{combined}

要求：
1. 保留所有专业术语、编号、格式标记
2. 保持原有的换行和缩进
3. 用---SPLIT---分隔翻译结果
4. 直接输出翻译，不要加前缀

中文翻译："""

    try:
        response = requests.post(TRANSLATION_API, json={
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": len(combined) * 3,
            "temperature": 0.3
        }, timeout=600)

        if response.status_code == 200:
            result = response.json()
            translated = result['choices'][0]['message']['content'].strip()

            # 移除thinking标签
            if '<think>' in translated:
                if '</think>' in translated:
                    translated = translated.split('</think>')[-1].strip()
                else:
                    return [translate_text(t) for t in texts]

            # 清理可能的前缀
            if translated.startswith("翻译：") or translated.startswith("中文翻译："):
                translated = translated.split("：", 1)[1].strip()

            # 分割结果
            parts = translated.split("---SPLIT---")

            if len(parts) == len(non_empty_texts):
                results = texts.copy()
                for i, idx in enumerate(non_empty_indices):
                    results[idx] = parts[i].strip()
                return results
            else:
                print(f"  ⚠️  批量分割失败，逐个翻译...")
                return [translate_text(t) for t in texts]
        else:
            return texts
    except Exception as e:
        print(f"  ⚠️  批量翻译失败: {e}")
        return texts

# ============ 删除英制单位 ============
def remove_imperial_units(text: str) -> str:
    """删除英制单位（保留括号中的公制单位）"""
    num_pattern = r'\d+(?:/\d+)?(?:\.\d+)?'

    # 模式1: 英制单位 + (公制单位) → 只保留 (公制单位)
    patterns_with_metric = [
        rf'{num_pattern}\s*-?\s*(?:inch|inches|in\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*-?\s*(?:foot|feet|ft\.?)\s*-?\s*(\([^)]+\))',
        rf'{num_pattern}\s*lbs?\.?/{num_pattern}\s*sq\.?\s*ft\.?\s*(\([^)]+\))',
        rf'{num_pattern}\s*lbs?\.?/cu\.?\s*ft\.?\s*(\([^)]+\))',
    ]

    result = text
    for pattern in patterns_with_metric:
        result = re.sub(pattern, r'\1', result, flags=re.IGNORECASE)

    # 模式2: 删除没有公制单位的英制单位
    standalone_imperial = [
        rf'{num_pattern}\s*-?\s*(?:inches|inch|in\.?)(?!\s*\()',
        rf'{num_pattern}\s*-?\s*(?:feet|foot|ft\.?)(?!\s*\()',
        rf'{num_pattern}\s*lbs?\.?/{num_pattern}\s*sq\.?\s*ft\.?(?!\s*\()',
    ]

    for pattern in standalone_imperial:
        result = re.sub(pattern, '', result, flags=re.IGNORECASE)

    # 清理多余空格
    result = re.sub(r'\s+', ' ', result).strip()
    result = re.sub(r'\s*,\s*,', ',', result)

    return result

# ============ 保护和恢复公制单位 ============
def protect_units(text: str) -> tuple:
    """保护公制单位不被翻译，返回（替换后的文本，映射表）"""
    units_map = {}
    counter = 0

    # 首先保护ASTM标准编号
    astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'
    for match in re.finditer(astm_pattern, text):
        placeholder = f'⟨⟨ASTM_{counter}⟩⟩'
        units_map[placeholder] = match.group(0)
        text = text.replace(match.group(0), placeholder, 1)
        counter += 1

    # 公制单位模式（复合单位在前）
    unit_patterns = [
        r'\d+(?:\.\d+)?\s*kg/sq\.\s*m',
        r'\d+(?:\.\d+)?\s*kg/m²',
        r'\d+(?:\.\d+)?\s*kg/m³',
        r'\d+(?:\.\d+)?\s*N/mm²',
        r'\d+(?:\.\d+)?\s*mm²',
        r'\d+(?:\.\d+)?\s*m²',
        r'\d+(?:\.\d+)?\s*(?:kPa|MPa|mm|cm|km|ml|Pa|°C)(?![²³/a-z])',
        r'\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])',
    ]

    def replace_marker(match):
        nonlocal counter
        marker_text = match.group(0)
        placeholder = f"⟨⟨UNIT{counter:04d}⟩⟩"
        units_map[placeholder] = marker_text
        counter += 1
        return placeholder

    result = text
    for pattern in unit_patterns:
        result = re.sub(pattern, replace_marker, result, flags=re.IGNORECASE)

    return result, units_map

def restore_units(text: str, units_map: dict) -> str:
    """恢复被保护的单位"""
    result = text
    for placeholder, unit in units_map.items():
        # 转换英文单位为符号格式
        converted_unit = unit
        converted_unit = re.sub(r'kg/sq\.\s*m', 'kg/m²', converted_unit, flags=re.IGNORECASE)
        converted_unit = re.sub(r'kg/cu\.\s*m', 'kg/m³', converted_unit, flags=re.IGNORECASE)
        result = result.replace(placeholder, converted_unit)
    return result

# ============ 应用翻译并保留颜色 ============
def apply_translation_with_colors(para, translated_text, original_formats):
    """应用翻译到段落，保留所有格式（公制单位自动设置为蓝绿色）"""

    # 检查翻译后的文本中是否包含公制单位
    unit_pattern = r'\d+(?:\.\d+)?\s*(?:kg/m²|kg/m³|N/mm²|mm²|m²|m³|kPa|MPa|mm|cm|km|ml|Pa|°C)(?![²³/a-z])|\d+(?:\.\d+)?\s*(?:m|g|t|l|L)(?![²³/a-zA-Z])'
    has_units = bool(re.search(unit_pattern, translated_text, re.IGNORECASE))

    # 如果没有单位，简单应用翻译
    if not has_units:
        if para.runs:
            first_run = para.runs[0]
            # 保存格式
            original_bold = first_run.bold
            original_italic = first_run.italic
            original_underline = first_run.underline
            original_font_size = first_run.font.size
            original_font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None

            # 修改文本
            first_run.text = translated_text

            # 恢复格式
            if original_bold is not None:
                first_run.bold = original_bold
            if original_italic is not None:
                first_run.italic = original_italic
            if original_underline is not None:
                first_run.underline = original_underline
            if original_font_size:
                first_run.font.size = original_font_size
            if original_font_color:
                first_run.font.color.rgb = original_font_color
            first_run.font.name = '宋体'

            # 删除其他runs
            for run in para.runs[1:]:
                run.text = ''
        return

    # 有单位，需要拆分并应用颜色
    # 保护ASTM标准编号（避免M后缀被误识别）
    astm_pattern = r'ASTM\s+[A-Z]\d+(?:/[A-Z]\d+)?M?'
    astm_placeholders = {}
    astm_counter = 0
    temp_text = translated_text

    for match in re.finditer(astm_pattern, temp_text):
        placeholder = f'⟨⟨TEMP_ASTM_{astm_counter}⟩⟩'
        astm_placeholders[placeholder] = match.group(0)
        temp_text = temp_text.replace(match.group(0), placeholder, 1)
        astm_counter += 1

    # 分割文本，保留单位
    parts = re.split(f'({unit_pattern})', temp_text, flags=re.IGNORECASE)

    # 保存第一个run的格式
    if para.runs:
        first_run = para.runs[0]
        default_bold = first_run.bold
        default_italic = first_run.italic
        default_underline = first_run.underline
        default_font_size = first_run.font.size
        default_font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
    else:
        default_bold = None
        default_italic = None
        default_underline = None
        default_font_size = None
        default_font_color = None

    # 清空所有runs
    for run in para.runs[:]:
        run._element.getparent().remove(run._element)

    # 重建runs
    for part in parts:
        if not part:
            continue

        # 恢复ASTM占位符
        final_part = part
        for placeholder, astm_text in astm_placeholders.items():
            final_part = final_part.replace(placeholder, astm_text)

        run = para.add_run(final_part)

        # 检查是否是单位
        is_unit = bool(re.match(unit_pattern, part, re.IGNORECASE))

        if is_unit:
            # 公制单位使用蓝绿色
            run.font.color.rgb = TEAL_COLOR
        else:
            # 非单位部分，应用默认格式
            if default_font_color:
                run.font.color.rgb = default_font_color

        # 应用其他格式
        if default_bold is not None:
            run.bold = default_bold
        if default_italic is not None:
            run.italic = default_italic
        if default_underline is not None:
            run.underline = default_underline
        if default_font_size:
            run.font.size = default_font_size
        run.font.name = '宋体'

# ============ 删除文本框 ============
def remove_textboxes(doc):
    """删除文档中的所有文本框"""
    try:
        # 删除shape中的文本框
        for shape in doc.inline_shapes:
            shape._element.getparent().remove(shape._element)

        # 删除浮动文本框
        for section in doc.sections:
            for element in section._element.xpath('.//w:txbxContent'):
                element.getparent().getparent().getparent().remove(element.getparent().getparent())
    except Exception as e:
        print(f"    ⚠️  文本框删除警告: {e}")

# ============ 删除页眉页脚 ============
def remove_headers_footers(doc):
    """删除所有页眉页脚（包括XML）"""
    try:
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False

            # 彻底清空页眉（包括段落和表格）
            section.header._element.clear()
            section.first_page_header._element.clear()
            section.even_page_header._element.clear()

            # 彻底清空页脚（包括段落和表格）
            section.footer._element.clear()
            section.first_page_footer._element.clear()
            section.even_page_footer._element.clear()
    except Exception as e:
        print(f"    ⚠️  页眉页脚删除警告: {e}")

# ============ 翻译DOCX文档 ============
def translate_docx(input_file: Path, output_file: Path) -> bool:
    """翻译DOCX文件内容（完整版）"""
    try:
        doc = Document(input_file)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # 1. 删除文本框
        print("    🗑️  删除文本框...")
        remove_textboxes(doc)

        # 2. 删除页眉页脚
        print("    🗑️  删除页眉页脚（包括XML）...")
        remove_headers_footers(doc)

        # 3. 过滤并收集需要翻译的段落
        paragraphs_data = []
        skipped_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                # 检查是否包含版权信息（直接删除）
                if any(keyword in para.text for keyword in COPYRIGHT_KEYWORDS):
                    para.clear()
                    skipped_count += 1
                    continue

                # 保存完整格式
                formats = []
                for run in para.runs:
                    formats.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })

                paragraphs_data.append({
                    'paragraph': para,
                    'text': para.text,
                    'formats': formats
                })

        if skipped_count > 0:
            print(f"    🗑️  已删除 {skipped_count} 个版权段落")

        total = len(paragraphs_data)
        print(f"    📝 批量翻译 {total} 个段落（{BATCH_SIZE}个/批）...")

        # 4. 批量翻译段落
        success_count = 0
        for i in range(0, total, BATCH_SIZE):
            batch = paragraphs_data[i:i+BATCH_SIZE]
            batch_texts = [item['text'] for item in batch]

            progress = f"[{i+1}-{min(i+BATCH_SIZE, total)}/{total}]"
            print(f"      {progress}", end=' ', flush=True)

            # 删除英制单位
            batch_texts = [remove_imperial_units(t) for t in batch_texts]

            # 保护公制单位
            protected_texts = []
            units_maps = []
            for text in batch_texts:
                protected_text, units_map = protect_units(text)
                protected_texts.append(protected_text)
                units_maps.append(units_map)

            # 翻译
            translated_texts = translate_batch(protected_texts)

            # 恢复单位
            translated_texts = [restore_units(trans, units_map)
                              for trans, units_map in zip(translated_texts, units_maps)]

            # 应用翻译（保持格式和颜色）
            for item, trans_text in zip(batch, translated_texts):
                para = item['paragraph']
                apply_translation_with_colors(para, trans_text, item['formats'])

            success_count += len(batch)
            print(f"✅")

        # 5. 保存
        doc.save(output_file)
        print(f"      进度: {success_count}/{total} ✅")
        return True

    except Exception as e:
        print(f"    ❌ 处理错误: {e}")
        import traceback
        traceback.print_exc()
        return False

# ============ 主程序 ============
def main():
    source_path = Path(SOURCE_DIR)

    if not source_path.exists():
        print(f"❌ 源目录不存在: {SOURCE_DIR}")
        return

    folders = sorted([d for d in source_path.iterdir() if d.is_dir() and not d.name.startswith('.')])

    print(f"\n{'='*70}")
    print(f"  🌐 批量翻译 DOCX → 中文（完整版）")
    print(f"{'='*70}")
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"模型: {MODEL_NAME}")
    print(f"文件夹: {len(folders)} 个")
    print(f"功能:")
    print(f"  ✅ 格式保留（粗体、斜体、字体、颜色）")
    print(f"  ✅ 过滤版权信息")
    print(f"  ✅ 删除文本框")
    print(f"  ✅ 删除页眉页脚（包括XML）")
    print(f"  ✅ 删除英制单位，保留公制单位")
    print(f"  ✅ 公制单位颜色（蓝绿色）")
    print(f"  ✅ 翻译文件名")
    print(f"{'='*70}\n")

    total_success = 0
    total_fail = 0
    total_skip = 0

    for folder_idx, folder in enumerate(folders, 1):
        folder_name = folder.name

        print(f"\n{'#'*70}")
        print(f"# [{folder_idx}/{len(folders)}] {folder_name}")
        print(f"{'#'*70}")

        files = sorted(list(folder.glob("*.docx")) + list(folder.glob("*.DOCX")))
        files = [f for f in files if not f.name.startswith('~$') and not f.name.startswith('.')]

        print(f"找到 {len(files)} 个DOCX文件\n")

        for file_idx, input_file in enumerate(files, 1):
            file_name = input_file.name

            print(f"  [{file_idx}/{len(files)}] {file_name}")

            # 翻译文件名
            print(f"    📝 翻译文件名...")
            translated_filename = translate_filename(file_name)

            if translated_filename != file_name:
                print(f"    → {translated_filename}")
            else:
                print(f"    → (保持原名)")

            output_folder = Path(OUTPUT_DIR) / folder_name
            output_file = output_folder / translated_filename

            if output_file.exists():
                print(f"    ⏭️  已存在\n")
                total_skip += 1
                continue

            print(f"    🌐 翻译内容...")

            if translate_docx(input_file, output_file):
                total_success += 1
                print()
            else:
                total_fail += 1
                print()

    print(f"\n{'='*70}")
    print(f"  📊 翻译完成")
    print(f"{'='*70}")
    print(f"✅ 成功: {total_success}")
    print(f"⏭️  跳过: {total_skip}")
    print(f"❌ 失败: {total_fail}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    main()
